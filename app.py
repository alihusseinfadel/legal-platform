import streamlit as st
import streamlit.components.v1 as components
import json
import datetime
import os
import base64
import io

st.set_page_config(
    page_title="منصة القانون الاداري - جامعة ديالى",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'About': "منصة القانون الاداري - جامعة ديالى\n\nتصميم: أ.م. علي حسين فاضل\nكادر وحدة تكنولوجيا المعلومات"
    }
)

# === PWA Meta Tags for mobile installation ===
st.markdown("""
<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0, user-scalable=yes">
<meta name="mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
<meta name="apple-mobile-web-app-title" content="القانون الاداري">
<meta name="application-name" content="منصة القانون الاداري">
<meta name="theme-color" content="#8B1A1A">
<meta name="description" content="منصة القانون الاداري - جامعة ديالى - استشارات قانونية ذكية">
<link rel="manifest" href="data:application/manifest+json;charset=utf-8;base64,ewogICJuYW1lIjogItmF2YbYtdipINin2YTZgtin2YbZiNmGINin2YTYp9iv2KfYsdmKIC0g2KzYp9mF2LnYqSDYr9mK2KfZhNmJIiwKICAic2hvcnRfbmFtZSI6ICLYp9mE2YLYp9mG2YjZhiDYp9mE2KfYr9in2LHZiiIsCiAgImRlc2NyaXB0aW9uIjogItmF2YbYtdipINin2YTYp9iz2KrYtNin2LHYp9iqINin2YTZgtin2YbZiNmG2YrYqSDYp9mE2KfYr9in2LHZitipIiwKICAic3RhcnRfdXJsIjogIi8iLAogICJkaXNwbGF5IjogInN0YW5kYWxvbmUiLAogICJvcmllbnRhdGlvbiI6ICJwb3J0cmFpdCIsCiAgImJhY2tncm91bmRfY29sb3IiOiAiIzhCMUExQSIsCiAgInRoZW1lX2NvbG9yIjogIiM4QjFBMUEiLAogICJsYW5nIjogImFyIiwKICAiZGlyIjogInJ0bCIKfQ==">
""", unsafe_allow_html=True)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800&display=swap');
    * { font-family: 'Tajawal', sans-serif !important; }

    .main .block-container { padding-top: 0.5rem; padding-bottom: 0.5rem; max-width: 1100px; position: relative; z-index: 1; }
    .main { background-color: #f5f5f5 !important; }
    [data-testid="stAppViewContainer"] { background-color: #f5f5f5; }
    /* Reduce vertical gaps */
    .main [data-testid="stVerticalBlock"] > div { gap: 0.4rem !important; }
    .main [data-testid="stHorizontalBlock"] { gap: 0.4rem !important; }
    .main hr { margin: 0.4rem 0 !important; }
    .main h1, .main h2, .main h3 { margin-top: 0.3rem !important; margin-bottom: 0.3rem !important; }
    .main [data-testid="element-container"] { margin: 0 !important; }

    [data-testid="stSidebar"] {
        direction: rtl; text-align: right;
        background: linear-gradient(180deg, #3d1f1f 0%, #2d1515 100%) !important;
        border-left: 1px solid #1a0808;
        box-shadow: -4px 0 25px rgba(0,0,0,0.3);
        min-width: 340px !important;
        width: 340px !important;
    }
    [data-testid="stSidebar"] > div:first-child {
        min-width: 340px !important;
        width: 340px !important;
        background: transparent !important;
    }
    /* Sidebar text colors - light on dark */
    [data-testid="stSidebar"] *,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] div {
        color: #f5e6d3 !important;
    }
    [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
        color: #fff !important;
    }
    [data-testid="stSidebar"] a { color: #d4a574 !important; }
    /* Hide ALL Material Icons / keyboard_arrow text that leaks through */
    [data-testid="stSidebar"] [data-testid="stIconMaterial"],
    [data-testid="stSidebar"] .material-icons,
    [data-testid="stSidebar"] .material-icons-outlined,
    [data-testid="stSidebar"] span[class*="icon"],
    .material-icons,
    .material-icons-outlined,
    [data-testid="stIconMaterial"],
    span[data-testid="stIconMaterial"] {
        font-family: 'Material Icons', 'Material Symbols Outlined', 'Material Symbols Rounded', sans-serif !important;
        font-size: 0 !important;
        width: 0 !important;
        height: 0 !important;
        overflow: hidden !important;
        display: inline-block !important;
    }
    /* Hide text content of Material Icons specifically */
    [data-testid="stIconMaterial"]::before,
    span.material-icons::before {
        content: attr(data-icon);
    }
    /* Hide password visibility toggle icon text */
    button[aria-label*="Show"] span,
    button[aria-label*="Hide"] span,
    button[title*="Show"] span,
    button[title*="Hide"] span {
        font-size: 0 !important;
    }
    button[aria-label*="Show password"]::before,
    button[title*="Show password"]::before {
        content: "👁"; font-size: 16px; color: #8B1A1A;
    }
    /* Hide sidebar collapse button icon */
    button[data-testid="stSidebarCollapseButton"] span,
    button[data-testid="baseButton-headerNoPadding"] span,
    [data-testid="stBaseButton-headerNoPadding"] span {
        font-size: 0 !important;
    }
    button[data-testid="stSidebarCollapseButton"]::before,
    button[data-testid="baseButton-headerNoPadding"]::before,
    [data-testid="stBaseButton-headerNoPadding"]::before {
        content: "☰"; font-size: 20px; color: #8B1A1A;
    }

    /* Dynamic font size based on user choice */
    .main p, .main li, .main label, .bot-bubble, .user-bubble,
    .msg-bubble, .law-card p, .consultation-card p {
        font-size: var(--user-font-size, 16px) !important;
    }

    /* ========== MOBILE RESPONSIVE (< 768px) ========== */
    @media screen and (max-width: 768px) {
        .main .block-container {
            padding: 0.5rem !important;
            max-width: 100% !important;
        }
        /* Sidebar collapses to drawer on mobile - Streamlit handles this */
        [data-testid="stSidebar"] {
            min-width: 85vw !important;
            width: 85vw !important;
        }
        [data-testid="stSidebar"] > div:first-child {
            min-width: 85vw !important;
            width: 85vw !important;
        }
        /* Hero + banners */
        .header-banner {
            padding: 1rem !important;
            margin-bottom: 1rem !important;
        }
        .header-banner h1 { font-size: 1.2rem !important; }
        .header-banner p { font-size: 0.85rem !important; }
        .page-title { font-size: 1.3rem !important; }
        /* Stat cards - wrap better */
        .stat-card {
            padding: 0.8rem !important;
            margin-bottom: 0.5rem;
        }
        .stat-card h2 { font-size: 1.5rem !important; }
        .stat-card p { font-size: 0.75rem !important; }
        /* Chat conversation full width */
        .chat-conversation {
            max-height: 400px !important;
            padding: 0.8rem 0.5rem !important;
        }
        .msg-content { max-width: 85% !important; }
        .msg-avatar { width: 34px !important; height: 34px !important; }
        .msg-bubble { padding: 0.7rem 0.9rem !important; font-size: 0.88rem !important; }
        /* Law cards */
        .law-card { padding: 0.8rem !important; }
        .law-card h4 { font-size: 0.95rem !important; }
        .law-card p { font-size: 0.85rem !important; }
        /* Consultation cards */
        .consultation-card { padding: 0.8rem !important; }
        /* Buttons full width on mobile */
        .stButton > button { padding: 0.5rem !important; font-size: 0.85rem !important; }
        /* Text inputs */
        .stTextInput input, .stTextArea textarea {
            font-size: 16px !important; /* Prevents iOS zoom */
        }
        /* Columns stack on mobile */
        [data-testid="stHorizontalBlock"] {
            flex-wrap: wrap !important;
            gap: 0.5rem !important;
        }
        [data-testid="stHorizontalBlock"] > div {
            min-width: 100% !important;
            flex: 1 1 100% !important;
        }
        /* Smaller watermark on mobile */
        [data-testid="stAppViewContainer"]::before {
            background-size: 120vw auto !important;
            opacity: 0.05 !important;
        }
        /* Tabs scroll horizontally */
        div[data-testid="stTabs"] [data-baseweb="tab-list"] {
            overflow-x: auto !important;
            flex-wrap: nowrap !important;
        }
    }

    /* ========== TABLET (768px - 1024px) ========== */
    @media screen and (min-width: 769px) and (max-width: 1024px) {
        .main .block-container {
            max-width: 95% !important;
            padding: 1rem !important;
        }
        [data-testid="stSidebar"] {
            min-width: 280px !important;
            width: 280px !important;
        }
    }
    /* Hide "Press Enter to apply" hint */
    [data-testid="InputInstructions"] { display: none !important; }
    /* Hide Streamlit branding */
    #MainMenu, footer, header[data-testid="stHeader"] { visibility: hidden; }
    /* Hide Deploy button */
    [data-testid="stToolbar"] { display: none !important; }
    [data-testid="stSidebar"] .stRadio > div {
        direction: rtl; gap: 2px !important;
    }
    [data-testid="stSidebar"] .stRadio > div > label {
        background: transparent;
        color: #e8d5bc !important;
        border-radius: 8px;
        padding: 0.7rem 1rem !important; margin: 0 !important;
        transition: all 0.25s ease;
        border: 1px solid transparent;
        font-weight: 600;
    }
    [data-testid="stSidebar"] .stRadio > div > label p {
        color: #e8d5bc !important;
    }
    [data-testid="stSidebar"] .stRadio > div > label:hover {
        background: rgba(255,255,255,0.06);
        border-color: rgba(212,165,116,0.2);
    }
    [data-testid="stSidebar"] .stRadio > div > label > div:first-child { display: none; }
    [data-testid="stSidebar"] .stRadio > div > label[data-baseweb="radio"]:has(input:checked),
    [data-testid="stSidebar"] .stRadio label:has(input:checked) {
        background: linear-gradient(135deg, #a0432a 0%, #8b3520 100%) !important;
        box-shadow: 0 4px 15px rgba(160,67,42,0.4), inset 0 1px 0 rgba(255,255,255,0.15) !important;
        border-color: #a0432a !important;
    }
    [data-testid="stSidebar"] .stRadio label:has(input:checked) p,
    [data-testid="stSidebar"] .stRadio label:has(input:checked) * {
        color: #fff !important;
    }

    /* === Sidebar Hero (Brand) === */
    .sb-hero {
        background: linear-gradient(135deg, rgba(255,255,255,0.08) 0%, rgba(255,255,255,0.02) 100%);
        border: 1px solid rgba(212,165,116,0.2);
        border-radius: 16px;
        padding: 1.3rem 1rem;
        text-align: center;
        margin: 0.2rem 0 1.2rem 0;
        position: relative;
        overflow: hidden;
    }
    .sb-hero::before {
        content: ""; position: absolute;
        top: -30px; right: -30px;
        width: 120px; height: 120px;
        background: rgba(212,165,116,0.08);
        border-radius: 50%;
    }
    .sb-hero::after {
        content: ""; position: absolute;
        bottom: -40px; left: -40px;
        width: 100px; height: 100px;
        background: rgba(212,165,116,0.08);
        border-radius: 50%;
    }
    .sb-hero-logo {
        width: 60px; height: 60px;
        background: linear-gradient(135deg, #d4a574 0%, #b8864e 100%);
        border-radius: 16px;
        display: flex; align-items: center; justify-content: center;
        color: #2d1515 !important; font-size: 1.8rem; font-weight: 800;
        margin: 0 auto 0.7rem auto;
        box-shadow: 0 6px 20px rgba(0,0,0,0.4);
        position: relative; z-index: 1;
    }
    .sb-hero h1 {
        color: #fff !important; margin: 0 !important;
        font-size: 1rem !important; font-weight: 800 !important;
        text-align: center !important;
        position: relative; z-index: 1;
    }
    .sb-hero p {
        color: #d4a574 !important; margin: 4px 0 0 0 !important;
        font-size: 0.72rem !important;
        text-align: center !important;
        font-weight: 600;
        position: relative; z-index: 1;
    }

    /* === Sidebar Divider === */
    .sb-divider {
        height: 1px;
        background: linear-gradient(90deg, transparent, rgba(212,165,116,0.25), transparent);
        margin: 1.2rem 0;
    }

    /* === Sidebar Labels (Section Titles) === */
    .sb-label {
        color: #d4a574 !important; font-size: 0.68rem; font-weight: 800;
        letter-spacing: 0.8px; text-transform: uppercase;
        margin: 1.3rem 0 0.6rem 0.3rem;
        direction: rtl; text-align: right;
        display: flex; align-items: center; gap: 8px;
        flex-direction: row-reverse;
    }
    .sb-label-bar {
        width: 3px; height: 14px; background: #d4a574;
        border-radius: 2px; display: inline-block;
    }

    /* === Status Card === */
    .sb-status {
        display: flex; align-items: center; gap: 10px;
        padding: 0.8rem; border-radius: 12px;
        margin-top: 0.5rem; direction: rtl;
    }
    .sb-status-on {
        background: rgba(34,197,94,0.12);
        border: 1px solid rgba(34,197,94,0.35);
    }
    .sb-status-off {
        background: rgba(245,158,11,0.12);
        border: 1px solid rgba(245,158,11,0.35);
    }
    .sb-status-icon {
        width: 32px; height: 32px; border-radius: 50%;
        display: flex; align-items: center; justify-content: center;
        color: #fff !important; font-size: 1.1rem; font-weight: 800;
        flex-shrink: 0;
    }
    .sb-status-on .sb-status-icon {
        background: #22c55e;
        box-shadow: 0 3px 10px rgba(34,197,94,0.4);
    }
    .sb-status-off .sb-status-icon {
        background: #f59e0b;
        box-shadow: 0 3px 10px rgba(245,158,11,0.4);
    }
    .sb-status-text { flex: 1; text-align: right; }
    .sb-status-t {
        color: #fff !important; font-size: 0.8rem; font-weight: 700;
    }
    .sb-status-d {
        color: #d4a574 !important; font-size: 0.68rem; margin-top: 1px;
    }

    /* === Stats Grid (2x2) === */
    .sb-stats-grid {
        display: grid; grid-template-columns: 1fr 1fr;
        gap: 8px;
    }
    .sb-stat {
        background: rgba(255,255,255,0.05);
        border: 1px solid rgba(212,165,116,0.15);
        border-radius: 10px; padding: 0.7rem 0.5rem;
        text-align: center;
        transition: all 0.25s;
    }
    .sb-stat:hover {
        border-color: #d4a574;
        transform: translateY(-2px);
        background: rgba(212,165,116,0.1);
    }
    .sb-stat-num {
        color: #d4a574 !important; font-size: 1.4rem; font-weight: 800;
        line-height: 1.1;
    }
    .sb-stat-lbl {
        color: #b8a07a !important; font-size: 0.7rem; font-weight: 600;
        margin-top: 3px;
    }

    /* === Credits Section === */
    .sb-credits {
        margin-top: 1.5rem;
        padding: 1rem 0.8rem;
        background: linear-gradient(135deg, rgba(212,165,116,0.08) 0%, rgba(212,165,116,0.02) 100%);
        border: 1px solid rgba(212,165,116,0.2);
        border-radius: 12px;
        direction: rtl;
        text-align: center;
    }
    .sb-credits-line {
        height: 2px;
        background: linear-gradient(90deg, transparent, #d4a574, transparent);
        margin-bottom: 0.8rem;
    }
    .sb-credits-label {
        color: #d4a574 !important;
        font-size: 0.7rem !important;
        font-weight: 700 !important;
        letter-spacing: 1px;
        margin: 0 0 0.6rem 0 !important;
        text-transform: uppercase;
    }
    .sb-credits-name {
        background: rgba(255,255,255,0.06);
        border: 1px solid rgba(212,165,116,0.25);
        border-radius: 10px;
        padding: 0.7rem 0.5rem;
        margin-bottom: 0.6rem;
    }
    .sb-credits-title {
        color: #fff !important;
        font-size: 0.9rem !important;
        font-weight: 800 !important;
        margin: 0 !important;
        line-height: 1.3;
    }
    .sb-credits-sub {
        color: #d4a574 !important;
        font-size: 0.7rem !important;
        margin: 3px 0 0 0 !important;
        font-weight: 600;
    }
    .sb-credits-team {
        display: flex;
        align-items: center;
        gap: 10px;
        padding: 0.6rem;
        background: rgba(255,255,255,0.04);
        border-radius: 8px;
        direction: rtl;
    }
    .sb-team-icon {
        width: 32px; height: 32px;
        background: linear-gradient(135deg, #d4a574, #b8864e);
        border-radius: 8px;
        display: flex; align-items: center; justify-content: center;
        font-size: 1rem;
        flex-shrink: 0;
    }
    .sb-team-text { flex: 1; text-align: right; }
    .sb-team-title {
        color: #fff !important;
        font-size: 0.75rem !important;
        font-weight: 700 !important;
        line-height: 1.3;
    }

    /* === Sidebar Footer === */
    .sb-footer {
        margin-top: 1rem; padding: 0.8rem 0.5rem 0.3rem 0.5rem;
        text-align: center; direction: rtl;
    }
    .sb-footer-line {
        height: 2px;
        background: linear-gradient(90deg, transparent, #d4a574, transparent);
        margin-bottom: 0.7rem;
    }
    .sb-footer-uni {
        color: #d4a574 !important; font-size: 0.8rem !important;
        font-weight: 800 !important;
        margin: 0 !important; text-align: center !important;
    }
    .sb-footer-col {
        color: #b8a07a !important; font-size: 0.7rem !important;
        margin: 2px 0 !important; text-align: center !important;
        font-weight: 600;
    }
    .sb-footer-year {
        color: #8b7355 !important; font-size: 0.65rem !important;
        margin: 5px 0 0 0 !important; text-align: center !important;
    }

    /* Sidebar input styling (dark theme) */
    [data-testid="stSidebar"] input[type="password"],
    [data-testid="stSidebar"] input[type="text"] {
        background: rgba(255,255,255,0.08) !important;
        color: #fff !important;
        border-radius: 10px !important;
        border: 1px solid rgba(212,165,116,0.25) !important;
        padding: 0.6rem !important;
        font-size: 0.85rem !important;
    }
    [data-testid="stSidebar"] input::placeholder { color: #8b7355 !important; }
    [data-testid="stSidebar"] input:focus {
        border-color: #d4a574 !important;
        box-shadow: 0 0 0 2px rgba(212,165,116,0.2) !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] {
        background: rgba(255,255,255,0.04) !important;
        border: 1px solid rgba(212,165,116,0.2) !important;
        border-radius: 10px !important;
        margin-top: 0.5rem;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary,
    [data-testid="stSidebar"] [data-testid="stExpander"] summary * {
        font-weight: 700; color: #d4a574 !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] [data-testid="stExpanderDetails"] * {
        color: #e8d5bc !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] strong {
        color: #fff !important;
    }
    /* Font size slider in sidebar */
    [data-testid="stSidebar"] [data-baseweb="slider"] * {
        color: #d4a574 !important;
    }
    [data-testid="stSidebar"] [data-baseweb="slider"] [role="slider"] {
        background: #d4a574 !important;
    }

    h1, h2, h3, h4, h5, h6, p, li, label, span, div { direction: rtl; text-align: right; }
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div { direction: rtl; text-align: right; font-family: 'Tajawal', sans-serif !important; }

    /* Top Header Bar */
    .top-header {
        background: #fff; border-bottom: 3px solid #8B1A1A;
        padding: 0.8rem 1.5rem; margin: -1rem -1rem 1.5rem -1rem;
        display: flex; justify-content: space-between; align-items: center;
        direction: rtl; border-radius: 0 0 8px 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
    }
    .top-header .logo-section { display: flex; align-items: center; gap: 12px; }
    .top-header .logo-section h2 { color: #333; margin: 0; font-size: 1.2rem; font-weight: 700; }
    .top-header .logo-icon {
        width: 42px; height: 42px; background: #8B1A1A; border-radius: 50%;
        display: flex; align-items: center; justify-content: center;
        color: #fff; font-size: 1.3rem; font-weight: 800;
    }

    /* Breadcrumb */
    .breadcrumb {
        color: #888; font-size: 0.85rem; margin-bottom: 0.5rem;
        direction: rtl; text-align: right;
    }
    .breadcrumb a { color: #8B1A1A; text-decoration: none; }

    /* Page Title */
    .page-title {
        font-size: 1.6rem; font-weight: 700; color: #333;
        margin-bottom: 1rem; direction: rtl; text-align: right;
        border-bottom: 3px solid #8B1A1A; padding-bottom: 0.5rem;
        display: inline-block;
    }

    /* Stat Cards */
    .stat-card {
        background: #fff; border: 1px solid #e0e0e0; border-radius: 10px;
        padding: 1.2rem; text-align: center; margin-bottom: 0.5rem;
        box-shadow: 0 1px 4px rgba(0,0,0,0.04); border-top: 3px solid #8B1A1A;
    }
    .stat-card h2 { text-align: center; color: #8B1A1A; margin: 0; font-size: 2rem; font-weight: 800; }
    .stat-card p { text-align: center; color: #666; margin: 0; font-size: 0.9rem; }

    /* Law Cards */
    .law-card {
        background: #fff; border: 1px solid #e0e0e0; border-radius: 10px;
        padding: 1.2rem; margin-bottom: 1rem; transition: border-color 0.3s;
        box-shadow: 0 1px 4px rgba(0,0,0,0.04); border-right: 4px solid #8B1A1A;
    }
    .law-card:hover { border-color: #8B1A1A; box-shadow: 0 2px 10px rgba(139,26,26,0.1); }
    .law-card h4 { color: #8B1A1A; margin-bottom: 0.5rem; text-align: right; }
    .law-card p { color: #555; font-size: 0.9rem; line-height: 1.8; text-align: right; }

    .law-tag {
        display: inline-block; background: #fdf0f0; color: #8B1A1A;
        padding: 2px 10px; border-radius: 20px; font-size: 0.75rem;
        margin-left: 5px; margin-top: 5px; border: 1px solid #e8c4c4;
    }

    /* Chat */
    .chat-user {
        background: #8B1A1A; border-radius: 12px 12px 0 12px; padding: 0.8rem 1.2rem;
        margin: 0.5rem 0; color: #fff; text-align: right; max-width: 80%;
        margin-right: auto; direction: rtl;
    }
    .chat-ai {
        background: #fff; border: 1px solid #e0e0e0; border-radius: 12px 12px 12px 0;
        padding: 0.8rem 1.2rem; margin: 0.5rem 0; color: #333; text-align: right;
        max-width: 80%; margin-left: auto; direction: rtl; line-height: 1.9;
        border-right: 3px solid #8B1A1A;
    }

    /* Status Badges */
    .status-badge { display: inline-block; padding: 3px 12px; border-radius: 20px; font-size: 0.8rem; font-weight: 600; }
    .status-new { background: #e8f5e9; color: #2e7d32; }
    .status-progress { background: #fff8e1; color: #f57f17; }
    .status-done { background: #e3f2fd; color: #1565c0; }
    .status-closed { background: #fce4ec; color: #c62828; }

    /* Consultation Card */
    .consultation-card {
        background: #fff; border: 1px solid #e0e0e0; border-radius: 10px;
        padding: 1.2rem; margin-bottom: 1rem;
        box-shadow: 0 1px 4px rgba(0,0,0,0.04); border-right: 4px solid #8B1A1A;
    }

    /* Header Banner */
    .header-banner {
        background: linear-gradient(135deg, #8B1A1A 0%, #a52a2a 50%, #6b1515 100%);
        border-radius: 12px; padding: 2rem; margin-bottom: 1.5rem;
        text-align: center; box-shadow: 0 4px 15px rgba(139,26,26,0.2);
    }
    .header-banner h1 { text-align: center; color: #fff; font-size: 1.8rem; margin-bottom: 0.3rem; }
    .header-banner p { text-align: center; color: #f0d0d0; font-size: 1rem; }

    /* Analysis Box */
    .analysis-box {
        background: #fff; border: 1px solid #8B1A1A; border-radius: 12px;
        padding: 1.5rem; margin: 1rem 0; line-height: 2.2;
        border-right: 4px solid #8B1A1A;
    }

    /* How it works box */
    .how-box {
        background: #fff; border: 1px solid #e0e0e0; border-radius: 10px;
        padding: 1rem; margin-bottom: 1rem; border-right: 4px solid #8B1A1A;
    }
    .how-box p { color: #555; }
    .how-box .how-title { color: #8B1A1A; font-weight: 700; margin: 0 0 0.5rem 0; }

    /* Tabs */
    div[data-testid="stTabs"] button { direction: rtl; }
    div[data-testid="stTabs"] button[aria-selected="true"] { border-bottom-color: #8B1A1A !important; color: #8B1A1A !important; }

    /* Buttons */
    .stButton > button { font-family: 'Tajawal', sans-serif !important; }
    .stButton > button[kind="primary"],
    .stButton > button[data-testid="stBaseButton-primary"] {
        background-color: #8B1A1A !important; border-color: #8B1A1A !important;
    }
    .stButton > button[data-testid="stBaseButton-primary"]:hover {
        background-color: #6b1515 !important; border-color: #6b1515 !important;
    }

    /* Warning/Info boxes */
    .warning-box {
        background: #fff8e1; border: 1px solid #ffe082; border-radius: 8px;
        padding: 0.8rem 1rem; direction: rtl;
    }
    .warning-box p { color: #e65100; margin: 0; font-size: 0.85rem; }

    /* File info */
    .file-info {
        background: #fff; border: 1px solid #e0e0e0; border-radius: 8px;
        padding: 0.8rem; margin: 0.5rem 0; border-right: 3px solid #8B1A1A;
    }
    .file-info p.name { color: #8B1A1A; margin: 0; font-weight: 600; }
    .file-info p.size { color: #888; margin: 0; font-size: 0.8rem; }

    /* Section cards */
    .section-card {
        background: #fff; border: 1px solid #e0e0e0; border-radius: 10px;
        padding: 1.2rem; text-align: center;
        box-shadow: 0 1px 4px rgba(0,0,0,0.04); transition: all 0.2s;
        border-top: 3px solid #8B1A1A;
    }
    .section-card:hover { box-shadow: 0 3px 12px rgba(139,26,26,0.12); }
    .section-card h2 { font-size: 1.1rem; margin-bottom: 0.3rem; text-align: center; }
    .section-card p { color: #666; font-size: 0.85rem; text-align: center; }

    /* === Modern Chatbot Interface === */

    /* Chat App Header (WhatsApp style) */
    .chat-app-header {
        background: linear-gradient(135deg, #8B1A1A 0%, #6b1515 100%);
        color: #fff; border-radius: 16px 16px 0 0;
        padding: 1rem 1.5rem; margin: 0.5rem 0 0 0;
        display: flex; justify-content: space-between; align-items: center;
        direction: rtl; box-shadow: 0 4px 20px rgba(139,26,26,0.25);
        position: relative;
    }
    .chat-header-center {
        display: flex; align-items: center; gap: 14px; flex: 1;
    }
    .bot-avatar-header {
        width: 56px; height: 56px;
        background: #fff; border-radius: 50%;
        display: flex; align-items: center; justify-content: center;
        color: #8B1A1A; font-size: 1.8rem; font-weight: 800;
        box-shadow: 0 4px 12px rgba(0,0,0,0.2);
        position: relative; flex-shrink: 0;
    }
    .online-dot {
        position: absolute; bottom: 2px; left: 2px;
        width: 14px; height: 14px;
        background: #22c55e; border-radius: 50%;
        border: 3px solid #fff;
    }
    .chat-header-info h3 {
        color: #fff !important; margin: 0; font-size: 1.25rem;
        text-align: right; font-weight: 700;
    }
    .chat-header-info p {
        color: #f0d0d0 !important; margin: 2px 0 0 0;
        font-size: 0.8rem; text-align: right;
    }
    .chat-header-info .status-text {
        font-weight: 600; margin-top: 4px;
    }
    .chat-header-left { display: flex; align-items: center; gap: 8px; }
    .bot-status-dot {
        width: 12px; height: 12px; border-radius: 50%;
        box-shadow: 0 0 0 4px rgba(255,255,255,0.2);
        animation: pulse-dot 2s infinite;
    }
    @keyframes pulse-dot {
        0%, 100% { box-shadow: 0 0 0 4px rgba(255,255,255,0.2); }
        50% { box-shadow: 0 0 0 8px rgba(255,255,255,0.1); }
    }

    /* Chat conversation area */
    .chat-conversation {
        background:
            linear-gradient(rgba(255,255,255,0.93), rgba(255,255,255,0.93)),
            repeating-linear-gradient(45deg, #fdf0f0 0, #fdf0f0 2px, #fff 2px, #fff 12px);
        border: 1px solid #e0e0e0;
        border-radius: 16px;
        padding: 1rem;
        max-height: 500px; overflow-y: auto;
        margin-bottom: 0.5rem;
        display: flex; flex-direction: column; gap: 0.5rem;
    }
    .chat-conversation .msg-row { margin-bottom: 0; }
    .chat-conversation::-webkit-scrollbar { width: 8px; }
    .chat-conversation::-webkit-scrollbar-track { background: #f5f5f5; }
    .chat-conversation::-webkit-scrollbar-thumb { background: #c0a0a0; border-radius: 4px; }

    /* Message Rows */
    .msg-row {
        display: flex; gap: 10px; margin-bottom: 1.2rem;
        direction: rtl; align-items: flex-end;
    }
    .msg-row.bot { justify-content: flex-start; }
    .msg-row.user { justify-content: flex-start; flex-direction: row-reverse; }

    .msg-avatar {
        width: 42px; height: 42px; flex-shrink: 0;
        border-radius: 50%;
        display: flex; align-items: center; justify-content: center;
        font-size: 1.1rem; font-weight: 800; color: #fff;
    }
    .bot-av {
        background: linear-gradient(135deg, #8B1A1A 0%, #a52a2a 100%);
        box-shadow: 0 3px 10px rgba(139,26,26,0.3);
    }
    .user-av {
        background: linear-gradient(135deg, #555 0%, #333 100%);
        box-shadow: 0 3px 10px rgba(0,0,0,0.2);
    }

    .msg-content { max-width: 75%; display: flex; flex-direction: column; }
    .user-content { align-items: flex-end; }

    .msg-bubble {
        padding: 0.9rem 1.2rem; line-height: 1.8;
        font-size: 0.95rem; word-wrap: break-word;
        direction: rtl; text-align: right;
    }
    .bot-bubble {
        background: #fff;
        border-radius: 18px 18px 18px 4px;
        color: #333; border: 1px solid #f0e0e0;
        box-shadow: 0 2px 8px rgba(139,26,26,0.06);
    }
    .bot-bubble p, .bot-bubble li, .bot-bubble strong, .bot-bubble h1, .bot-bubble h2, .bot-bubble h3, .bot-bubble h4 {
        text-align: right !important; direction: rtl;
    }
    .user-bubble {
        background: linear-gradient(135deg, #8B1A1A 0%, #a52a2a 100%);
        color: #fff !important;
        border-radius: 18px 18px 4px 18px;
        box-shadow: 0 2px 10px rgba(139,26,26,0.25);
    }
    .user-bubble * { color: #fff !important; }

    .msg-time {
        font-size: 0.7rem; color: #999;
        margin: 4px 8px 0 8px; direction: rtl; text-align: right;
        font-weight: 600;
    }
    .user-time { text-align: left; color: #8B1A1A; }

    /* Suggestions */
    .suggestions-title {
        color: #8B1A1A; font-size: 0.95rem; font-weight: 700;
        margin: 0.8rem 0 0.5rem 0; direction: rtl; text-align: right;
    }

    /* Message Input Bar */
    .message-input-bar {
        background: #fff; border: 2px solid #8B1A1A;
        border-radius: 30px; padding: 0.4rem 0.8rem;
        margin-top: 0.8rem;
        box-shadow: 0 4px 15px rgba(139,26,26,0.1);
    }
    .message-input-bar .stTextInput > div > div > input {
        border: none !important; background: transparent !important;
        font-size: 1rem; padding: 0.5rem !important;
    }
    .message-input-bar .stTextInput > div { border: none !important; }
    .message-input-bar .stButton > button {
        border-radius: 50% !important;
        font-size: 1.2rem !important;
        height: 42px !important; min-width: 42px !important;
        padding: 0 !important;
    }
</style>
""", unsafe_allow_html=True)

# === Data ===
APP_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(APP_DIR, "data")
DATA_FILE = os.path.join(DATA_DIR, "data.json")
GEO_CONFIG_FILE = os.path.join(DATA_DIR, "geo_config.json")
GEO_LOG_FILE = os.path.join(DATA_DIR, "geo_log.json")
CONVERSATIONS_FILE = os.path.join(DATA_DIR, "conversations.json")
os.makedirs(DATA_DIR, exist_ok=True)

def load_conversations():
    if os.path.exists(CONVERSATIONS_FILE):
        try:
            with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return []

def save_conversations(convs):
    with open(CONVERSATIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(convs, f, ensure_ascii=False, indent=2)

def new_conversation_id():
    return f"CONV-{datetime.datetime.now().strftime('%Y%m%d%H%M%S%f')[:20]}"

# === Predefined University of Diyala Buildings ===
UOD_BUILDINGS = [
    {"id": 1, "name": "رئاسة جامعة ديالى", "lat": 33.763200, "lon": 44.618300, "type": "رئاسة"},
    {"id": 2, "name": "كلية القانون والعلوم السياسية", "lat": 33.762500, "lon": 44.617100, "type": "كلية"},
    {"id": 3, "name": "كلية التربية للعلوم الصرفة", "lat": 33.764100, "lon": 44.619500, "type": "كلية"},
    {"id": 4, "name": "كلية التربية للعلوم الانسانية", "lat": 33.763900, "lon": 44.620300, "type": "كلية"},
    {"id": 5, "name": "كلية العلوم", "lat": 33.764800, "lon": 44.618900, "type": "كلية"},
    {"id": 6, "name": "كلية الهندسة", "lat": 33.762100, "lon": 44.620100, "type": "كلية"},
    {"id": 7, "name": "كلية الطب", "lat": 33.765500, "lon": 44.619200, "type": "كلية"},
    {"id": 8, "name": "كلية الزراعة", "lat": 33.761400, "lon": 44.621800, "type": "كلية"},
    {"id": 9, "name": "كلية الطب البيطري", "lat": 33.761800, "lon": 44.622500, "type": "كلية"},
    {"id": 10, "name": "كلية الادارة والاقتصاد", "lat": 33.763500, "lon": 44.616800, "type": "كلية"},
    {"id": 11, "name": "كلية التربية البدنية وعلوم الرياضة", "lat": 33.762800, "lon": 44.622900, "type": "كلية"},
    {"id": 12, "name": "كلية العلوم الاسلامية", "lat": 33.763000, "lon": 44.616500, "type": "كلية"},
    {"id": 13, "name": "كلية الفنون الجميلة", "lat": 33.764300, "lon": 44.617600, "type": "كلية"},
    {"id": 14, "name": "كلية التربية الاساسية", "lat": 33.765100, "lon": 44.620500, "type": "كلية"},
    {"id": 15, "name": "كلية الهندسة - قسم الحاسوب", "lat": 33.762000, "lon": 44.619800, "type": "قسم"},
    {"id": 16, "name": "مركز الحاسبة الالكترونية", "lat": 33.763400, "lon": 44.618800, "type": "مركز"},
    {"id": 17, "name": "المكتبة المركزية", "lat": 33.763700, "lon": 44.618500, "type": "مبنى اداري"},
    {"id": 18, "name": "قاعة الاحتفالات الكبرى", "lat": 33.763300, "lon": 44.619000, "type": "مبنى اداري"},
]

# Default: University of Diyala main campus (Baqubah)
DEFAULT_GEO_CONFIG = {
    "markers": [],
    "polygon": [],  # List of [lat, lon] points forming the allowed polygon
    "admin_password": "admin123"
}

def load_data():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"consultations": [], "analyses": []}

def save_data(data):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_geo_config():
    if os.path.exists(GEO_CONFIG_FILE):
        try:
            with open(GEO_CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return dict(DEFAULT_GEO_CONFIG)

def save_geo_config(cfg):
    with open(GEO_CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)

def load_geo_log():
    if os.path.exists(GEO_LOG_FILE):
        try:
            with open(GEO_LOG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return []

def log_geo_attempt(lat, lon, allowed, zone_id=None, distance_km=None):
    log = load_geo_log()
    log.append({
        "timestamp": datetime.datetime.now().isoformat(),
        "lat": lat, "lon": lon,
        "allowed": allowed,
        "zone_id": zone_id,
        "distance_km": distance_km
    })
    # Keep only last 10000 entries
    log = log[-10000:]
    with open(GEO_LOG_FILE, "w", encoding="utf-8") as f:
        json.dump(log, f, ensure_ascii=False, indent=2)

def haversine_km(lat1, lon1, lat2, lon2):
    """Calculate great-circle distance in kilometers using Haversine formula"""
    import math
    R = 6371.0
    dlat = math.radians(lat2 - lat1)
    dlon = math.radians(lon2 - lon1)
    a = (math.sin(dlat/2)**2 +
         math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) *
         math.sin(dlon/2)**2)
    return 2 * R * math.asin(math.sqrt(a))

def point_in_polygon(lat, lon, polygon):
    """Ray casting algorithm for point-in-polygon. polygon is list of [lat, lon]"""
    if not polygon or len(polygon) < 3:
        return False
    n = len(polygon)
    inside = False
    j = n - 1
    for i in range(n):
        lat_i, lon_i = polygon[i][0], polygon[i][1]
        lat_j, lon_j = polygon[j][0], polygon[j][1]
        if ((lon_i > lon) != (lon_j > lon)) and \
           (lat < (lat_j - lat_i) * (lon - lon_i) / (lon_j - lon_i + 1e-12) + lat_i):
            inside = not inside
        j = i
    return inside

def nearest_polygon_distance_km(lat, lon, polygon):
    """Min distance from point to any polygon vertex (approximation)"""
    if not polygon:
        return float('inf')
    return min(haversine_km(lat, lon, p[0], p[1]) for p in polygon)

def convex_hull(points):
    """Compute convex hull of list of [lat, lon] using Andrew's monotone chain"""
    pts = sorted(set(tuple(p) for p in points))
    if len(pts) <= 1:
        return [list(p) for p in pts]
    def cross(o, a, b):
        return (a[0]-o[0]) * (b[1]-o[1]) - (a[1]-o[1]) * (b[0]-o[0])
    lower = []
    for p in pts:
        while len(lower) >= 2 and cross(lower[-2], lower[-1], p) <= 0:
            lower.pop()
        lower.append(p)
    upper = []
    for p in reversed(pts):
        while len(upper) >= 2 and cross(upper[-2], upper[-1], p) <= 0:
            upper.pop()
        upper.append(p)
    return [list(p) for p in lower[:-1] + upper[:-1]]

def check_in_zones(lat, lon, geo_config):
    """Check if point is within the polygon. Returns (allowed, distance_to_nearest_vertex)"""
    polygon = geo_config.get("polygon", [])
    if polygon and len(polygon) >= 3:
        inside = point_in_polygon(lat, lon, polygon)
        if inside:
            return True, 0.0
        else:
            return False, nearest_polygon_distance_km(lat, lon, polygon)
    # Fallback: if no polygon, check distance from main marker within 3km
    markers = geo_config.get("markers", [])
    main = next((m for m in markers if m.get("is_main")), None) or (markers[0] if markers else None)
    if main:
        d = haversine_km(lat, lon, main["lat"], main["lon"])
        return (d <= 3.0), d
    return False, float('inf')

if "data" not in st.session_state:
    st.session_state.data = load_data()
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = []
if "analysis_result" not in st.session_state:
    st.session_state.analysis_result = None
if "geo_verified" not in st.session_state:
    st.session_state.geo_verified = True  # Direct access mode (no geo check)
if "geo_status" not in st.session_state:
    st.session_state.geo_status = "direct"
if "geo_config" not in st.session_state:
    st.session_state.geo_config = load_geo_config()
if "admin_logged_in" not in st.session_state:
    st.session_state.admin_logged_in = False
if "font_size" not in st.session_state:
    st.session_state.font_size = "متوسط"
if "conversations" not in st.session_state:
    # Start fresh - no previous conversations
    st.session_state.conversations = []
if "current_conv_id" not in st.session_state:
    st.session_state.current_conv_id = None
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = []

# Check URL query params for geo verification response
try:
    qp = st.query_params
    if "geo_ok" in qp and "lat" in qp and "lon" in qp:
        lat = float(qp["lat"])
        lon = float(qp["lon"])
        allowed, dist = check_in_zones(lat, lon, st.session_state.geo_config)
        if allowed:
            st.session_state.geo_verified = True
            st.session_state.geo_status = "verified"
            log_geo_attempt(lat, lon, True, None, round(dist, 3))
        else:
            st.session_state.geo_status = "far"
            log_geo_attempt(lat, lon, False, None, round(dist, 3))
        st.query_params.clear()
    elif "geo_fail" in qp:
        st.session_state.geo_status = "denied"
        st.query_params.clear()
    elif "admin_bypass" in qp:
        # Emergency bypass for admin only
        st.session_state.geo_verified = True
        st.session_state.geo_status = "bypass"
        st.query_params.clear()
    elif "save_config" in qp and st.session_state.admin_logged_in:
        try:
            import urllib.parse
            cfg_json = urllib.parse.unquote(qp["save_config"])
            new_cfg = json.loads(cfg_json)
            # Preserve admin password
            new_cfg["admin_password"] = st.session_state.geo_config.get("admin_password", "admin123")
            st.session_state.geo_config = new_cfg
            save_geo_config(new_cfg)
        except Exception as e:
            st.session_state.save_error = str(e)
        st.query_params.clear()
    elif "add_pt" in qp and st.session_state.admin_logged_in:
        try:
            parts = qp["add_pt"].split(",")
            if len(parts) == 2:
                if "polygon_draft" not in st.session_state:
                    st.session_state.polygon_draft = []
                st.session_state.polygon_draft.append([float(parts[0]), float(parts[1])])
        except Exception:
            pass
        st.query_params.clear()
except Exception:
    pass

# === Apply dynamic font size ===
_font_size_map = {"صغير": "14px", "متوسط": "16px", "كبير": "18px", "كبير جداً": "20px"}
_user_font = _font_size_map.get(st.session_state.get("font_size", "متوسط"), "16px")
st.markdown(f'<style>:root {{ --user-font-size: {_user_font}; }}</style>', unsafe_allow_html=True)

# === Load University seal logo for watermark (local file, cached) ===
@st.cache_data
def get_diyala_logo_b64():
    try:
        logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "diyala_full_seal.png")
        if os.path.exists(logo_path):
            with open(logo_path, "rb") as f:
                return base64.b64encode(f.read()).decode("ascii")
        # Fallback: download if local file missing
        import urllib.request, urllib.parse
        url = "https://uodiyala.edu.iq/wp-content/uploads/2023/04/" + urllib.parse.quote("شعار-الجامعة-الجديد-3") + ".png"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            return base64.b64encode(resp.read()).decode("ascii")
    except Exception:
        return None

DIYALA_LOGO_B64 = get_diyala_logo_b64()

# === Geo Verification Gate ===
if not st.session_state.geo_verified:
    st.markdown("""
    <style>
    [data-testid="stSidebar"] { display: none !important; }
    .main .block-container { max-width: 700px !important; padding-top: 3rem !important; }
    header[data-testid="stHeader"] { display: none; }
    </style>
    """, unsafe_allow_html=True)

    if DIYALA_LOGO_B64:
        st.markdown(f"""
        <div style="text-align:center;margin-bottom:1rem;">
            <img src="data:image/png;base64,{DIYALA_LOGO_B64}" style="width:160px;height:auto;">
        </div>
        """, unsafe_allow_html=True)

    st.markdown("""
    <div style="text-align:center;margin-bottom:1.5rem;">
        <h1 style="color:#8B1A1A;margin:0.5rem 0;font-size:1.6rem;">منصة القانون الاداري</h1>
        <p style="color:#666;margin:0;font-size:1rem;">جامعة ديالى - رئاسة الجامعة</p>
    </div>
    """, unsafe_allow_html=True)

    status = st.session_state.geo_status
    polygon = st.session_state.geo_config.get("polygon", [])
    markers = st.session_state.geo_config.get("markers", [])

    if status == "denied":
        st.error("تم رفض الوصول للموقع. هذه المنصة تتطلب التحقق من الموقع الجغرافي.")
    elif status == "far":
        st.error("موقعك خارج حدود المناطق المسموحة لاستخدام هذه المنصة.")
    else:
        st.info("للوصول الى المنصة، يرجى السماح بمشاركة موقعك الجغرافي للتحقق من تواجدك ضمن حدود جامعة ديالى.")

    # Send polygon to JavaScript
    polygon_js = json.dumps(polygon)
    markers_js = json.dumps([{"lat": m["lat"], "lon": m["lon"]} for m in markers if m.get("is_main")])

    # Geolocation verification component
    components.html(f"""
    <!DOCTYPE html>
    <html><head><meta charset="utf-8"></head>
    <body style="margin:0;padding:0;background:transparent;font-family:'Tajawal',sans-serif;direction:rtl;">
    <div style="text-align:center;padding:1rem;">
        <button id="verifyBtn" style="
            background:linear-gradient(135deg,#8B1A1A 0%,#a52a2a 100%);
            color:#fff;border:none;padding:14px 40px;border-radius:30px;
            font-size:1.05rem;font-weight:700;cursor:pointer;font-family:'Tajawal',sans-serif;
            box-shadow:0 6px 20px rgba(139,26,26,0.3);">
            التحقق من الموقع والدخول
        </button>
        <p id="verifyStatus" style="color:#888;font-size:0.85rem;margin-top:12px;">اضغط الزر وسيطلب المتصفح اذن مشاركة الموقع</p>
    </div>
    <script>
    (function() {{
        const btn = document.getElementById('verifyBtn');
        const statusEl = document.getElementById('verifyStatus');
        const POLYGON = {polygon_js};
        const MAIN_MARKERS = {markers_js};

        function haversine(lat1, lon1, lat2, lon2) {{
            const R = 6371;
            const dLat = (lat2 - lat1) * Math.PI / 180;
            const dLon = (lon2 - lon1) * Math.PI / 180;
            const a = Math.sin(dLat/2) ** 2 +
                      Math.cos(lat1 * Math.PI/180) * Math.cos(lat2 * Math.PI/180) *
                      Math.sin(dLon/2) ** 2;
            return 2 * R * Math.asin(Math.sqrt(a));
        }}

        function pointInPolygon(lat, lon, poly) {{
            if (!poly || poly.length < 3) return false;
            let inside = false;
            let j = poly.length - 1;
            for (let i = 0; i < poly.length; i++) {{
                const latI = poly[i][0], lonI = poly[i][1];
                const latJ = poly[j][0], lonJ = poly[j][1];
                if (((lonI > lon) !== (lonJ > lon)) &&
                    (lat < (latJ - latI) * (lon - lonI) / (lonJ - lonI + 1e-12) + latI)) {{
                    inside = !inside;
                }}
                j = i;
            }}
            return inside;
        }}

        function sendResult(params) {{
            const parent = window.parent;
            const url = new URL(parent.location.href);
            for (const k in params) url.searchParams.set(k, params[k]);
            parent.location.href = url.toString();
        }}

        btn.addEventListener('click', function() {{
            if (!('geolocation' in navigator)) {{
                statusEl.textContent = 'المتصفح لا يدعم تحديد الموقع';
                statusEl.style.color = '#c62828';
                return;
            }}
            statusEl.textContent = 'جاري تحديد موقعك...';
            statusEl.style.color = '#8B1A1A';
            btn.disabled = true;

            navigator.geolocation.getCurrentPosition(
                function(pos) {{
                    const lat = pos.coords.latitude;
                    const lon = pos.coords.longitude;
                    let allowed = false;
                    let minDist = Infinity;

                    if (POLYGON.length >= 3) {{
                        allowed = pointInPolygon(lat, lon, POLYGON);
                        for (const p of POLYGON) {{
                            const d = haversine(lat, lon, p[0], p[1]);
                            if (d < minDist) minDist = d;
                        }}
                    }} else if (MAIN_MARKERS.length > 0) {{
                        // Fallback: use main marker with 3km radius
                        for (const m of MAIN_MARKERS) {{
                            const d = haversine(lat, lon, m.lat, m.lon);
                            if (d < minDist) minDist = d;
                        }}
                        allowed = minDist <= 3.0;
                    }}

                    // Always send coordinates for logging
                    setTimeout(function() {{
                        sendResult({{geo_ok: '1', lat: lat.toFixed(6), lon: lon.toFixed(6)}});
                    }}, 800);

                    if (allowed) {{
                        statusEl.textContent = 'تم التحقق بنجاح! جاري فتح المنصة...';
                        statusEl.style.color = '#2e7d32';
                    }} else {{
                        statusEl.textContent = 'موقعك خارج الحدود (' + minDist.toFixed(1) + ' كم من اقرب نقطة)';
                        statusEl.style.color = '#c62828';
                    }}
                }},
                function(err) {{
                    statusEl.textContent = 'تم رفض اذن الموقع: ' + err.message;
                    statusEl.style.color = '#c62828';
                    btn.disabled = false;
                    setTimeout(function() {{ sendResult({{geo_fail: '1'}}); }}, 1500);
                }},
                {{ enableHighAccuracy: true, timeout: 15000, maximumAge: 0 }}
            );
        }});
    }})();
    </script>
    </body></html>
    """, height=140)

    polygon_info = f"مضلع محدد بـ <b>{len(polygon)}</b> نقطة" if len(polygon) >= 3 else "لم يُرسم المضلع بعد"
    st.markdown(f"""
    <div style="background:#fdf0f0;border:1px solid #e8c4c4;border-radius:10px;padding:1rem;margin-top:1rem;">
        <p style="color:#8B1A1A;margin:0 0 0.5rem 0;font-weight:700;">ملاحظة:</p>
        <p style="color:#555;margin:0;font-size:0.9rem;line-height:1.8;">
            هذه المنصة مخصصة حصرياً لطلاب وتدريسيي جامعة ديالى<br>
            الحدود الجغرافية المحددة: {polygon_info}<br>
            يتم حفظ سجل الدخول للاحصائيات فقط
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Hidden admin link
    with st.expander("دخول الادمن"):
        pwd = st.text_input("كلمة المرور", type="password", key="admin_pwd_gate")
        if st.button("دخول", key="admin_btn_gate"):
            if pwd == st.session_state.geo_config.get("admin_password", "admin123"):
                st.session_state.admin_logged_in = True
                st.session_state.geo_verified = True
                st.session_state.geo_status = "admin"
                st.rerun()
            else:
                st.error("كلمة مرور خاطئة")

    st.stop()

# === Watermark with full University seal (covers entire page) ===
if DIYALA_LOGO_B64:
    st.markdown(
        '<style>'
        '[data-testid="stAppViewContainer"]::before {'
        'content: ""; position: fixed; top: 0; left: 0; right: 0; bottom: 0;'
        'background-image: url("data:image/png;base64,' + DIYALA_LOGO_B64 + '");'
        'background-repeat: no-repeat;'
        'background-position: center center;'
        'background-size: min(95vh, 95vw) auto;'
        'opacity: 0.08;'
        'pointer-events: none; z-index: 0;'
        '}'
        '.main .block-container { position: relative; z-index: 1; }'
        '[data-testid="stSidebar"] { z-index: 2; background-color: rgba(255,255,255,0.97) !important; }'
        '</style>',
        unsafe_allow_html=True
    )

# === Laws DB ===
LAWS_DB = [
    {"id": 1, "title": "قانون انضباط موظفي الدولة والقطاع العام رقم (14) لسنة 1991", "category": "انضباط وظيفي",
     "summary": "ينظم الاحكام المتعلقة بانضباط موظفي الدولة والقطاع العام ويحدد العقوبات الانضباطية من لفت النظر الى العزل.",
     "tags": ["انضباط", "عقوبات", "موظفين", "قطاع عام"],
     "articles": [
         {"num": "المادة 1", "text": "تسري احكام هذا القانون على موظفي الدولة والقطاع العام"},
         {"num": "المادة 4", "text": "لا يجوز فرض عقوبة انضباطية على الموظف الا بعد اجراء تحقيق معه"},
         {"num": "المادة 8", "text": "العقوبات الانضباطية: لفت النظر، الانذار، قطع الراتب، التوبيخ، انقاص الراتب، تنزيل الدرجة، الفصل، العزل"},
         {"num": "المادة 10", "text": "للموظف حق الاعتراض على العقوبة المفروضة عليه خلال ثلاثين يوما"},
         {"num": "المادة 15", "text": "تشكل لجنة انضباطية في كل دائرة برئاسة موظف لا تقل درجته عن مدير"},
     ]},
    {"id": 2, "title": "قانون الخدمة المدنية رقم (24) لسنة 1960", "category": "خدمة مدنية",
     "summary": "ينظم شؤون التوظيف من حيث التعيين والترقية والنقل والاجازات والرواتب وانتهاء الخدمة.",
     "tags": ["تعيين", "ترقية", "رواتب", "اجازات", "خدمة مدنية"],
     "articles": [
         {"num": "المادة 2", "text": "الوظائف الحكومية حق للعراقيين ولا يجوز حرمان احد منها الا وفق القانون"},
         {"num": "المادة 7", "text": "يشترط في المرشح للتعيين ان يكون عراقي الجنسية واكمل الثامنة عشرة"},
         {"num": "المادة 12", "text": "تكون مدة التجربة سنة واحدة من تاريخ مباشرة الموظف"},
         {"num": "المادة 29", "text": "يستحق الموظف اجازة اعتيادية براتب تام مدتها ثلاثون يوما في السنة"},
         {"num": "المادة 36", "text": "تنتهي خدمة الموظف بالاستقالة او الاحالة على التقاعد او الفصل او العزل"},
     ]},
    {"id": 3, "title": "قانون التقاعد الموحد رقم (9) لسنة 2014", "category": "تقاعد",
     "summary": "ينظم حقوق التقاعد ويوحد الانظمة التقاعدية ويحدد شروط استحقاق الراتب التقاعدي.",
     "tags": ["تقاعد", "راتب تقاعدي", "مكافأة", "خدمة"],
     "articles": [
         {"num": "المادة 9", "text": "يستحق الموظف راتبا تقاعديا اذا كانت لديه خدمة لا تقل عن 15 سنة"},
         {"num": "المادة 10", "text": "يحتسب الراتب التقاعدي على اساس 2.5% من معدل الراتب عن كل سنة خدمة"},
         {"num": "المادة 21", "text": "يستحق ذوو المتقاعد المتوفى راتبا تقاعديا وفق الحصص المقررة"},
     ]},
    {"id": 4, "title": "قانون الملاك رقم (25) لسنة 1960", "category": "ملاك وظيفي",
     "summary": "ينظم الهيكل الوظيفي والملاك من حيث الدرجات والعناوين والرواتب والعلاوات.",
     "tags": ["ملاك", "درجات", "رواتب", "علاوات"],
     "articles": [
         {"num": "المادة 1", "text": "يراد بالملاك جدول تنظيم يبين عدد ونوع الوظائف الدائمة"},
         {"num": "المادة 3", "text": "تقسم الوظائف الى عشر درجات"},
         {"num": "المادة 5", "text": "يمنح الموظف علاوة سنوية وفق الجدول الملحق"},
     ]},
    {"id": 5, "title": "قانون حقوق والتزامات الموظف رقم (22) لسنة 2008", "category": "حقوق الموظف",
     "summary": "يحدد حقوق الموظف والتزاماته تجاه الدولة والمواطنين.",
     "tags": ["حقوق", "التزامات", "موظف", "نزاهة"],
     "articles": [
         {"num": "المادة 2", "text": "للموظف الحق في الحصول على راتبه ومخصصاته في مواعيدها"},
         {"num": "المادة 4", "text": "يلتزم الموظف باداء واجبات وظيفته بدقة وامانة"},
         {"num": "المادة 7", "text": "يحظر على الموظف استغلال وظيفته لتحقيق منفعة شخصية"},
     ]},
    {"id": 6, "title": "قانون مجلس شورى الدولة رقم (65) لسنة 1979", "category": "قضاء اداري",
     "summary": "ينظم عمل مجلس شورى الدولة كجهة قضائية ادارية عليا للطعون والفتاوى.",
     "tags": ["قضاء اداري", "طعون", "فتاوى", "شورى"],
     "articles": [
         {"num": "المادة 1", "text": "مجلس شورى الدولة هيئة قضائية ادارية مستقلة"},
         {"num": "المادة 7", "text": "يختص المجلس بالنظر في صحة الاوامر والقرارات الادارية"},
         {"num": "المادة 12", "text": "للمتضرر الطعن بالقرار الاداري خلال ستين يوما من تاريخ تبليغه"},
     ]},
    {"id": 7, "title": "قانون رواتب موظفي الدولة رقم (22) لسنة 2008", "category": "رواتب",
     "summary": "يحدد سلم رواتب موظفي الدولة والمخصصات والعلاوات.",
     "tags": ["رواتب", "مخصصات", "علاوات", "سلم رواتب"],
     "articles": [
         {"num": "المادة 2", "text": "يتكون الراتب الكلي من الراتب الاسمي مضافا اليه المخصصات"},
         {"num": "المادة 3", "text": "تحدد الدرجات الوظيفية من الاولى الى العاشرة"},
         {"num": "المادة 8", "text": "يستحق الموظف مخصصات شهادة حسب التحصيل الدراسي"},
     ]},
    {"id": 8, "title": "قانون العقوبات العراقي رقم (111) لسنة 1969 - الجرائم الوظيفية", "category": "جرائم وظيفية",
     "summary": "الاحكام الجزائية للجرائم الوظيفية: الرشوة والاختلاس والتزوير.",
     "tags": ["رشوة", "اختلاس", "تزوير", "جرائم وظيفية"],
     "articles": [
         {"num": "المادة 307", "text": "يعاقب بالسجن كل موظف قبل او طلب هدية لاداء عمل وظيفي"},
         {"num": "المادة 315", "text": "يعاقب بالسجن كل موظف اختلس او استولى على مال عام"},
         {"num": "المادة 331", "text": "يعاقب بالحبس كل موظف استغل وظيفته فاضر بمصلحة الدولة"},
     ]},
    {"id": 9, "title": "قانون المرافعات المدنية رقم (83) لسنة 1969", "category": "مرافعات",
     "summary": "ينظم اجراءات التقاضي امام المحاكم المدنية والاستئناف والتمييز.",
     "tags": ["مرافعات", "دعوى", "استئناف", "تمييز"],
     "articles": [
         {"num": "المادة 2", "text": "لا يقبل اي طلب لا تكون لصاحبه فيه مصلحة قائمة"},
         {"num": "المادة 30", "text": "ترفع الدعوى بعريضة تودع في قلم كتاب المحكمة"},
         {"num": "المادة 187", "text": "يجوز الطعن تمييزا في الاحكام الصادرة من محاكم الاستئناف"},
     ]},
    {"id": 10, "title": "قانون التعليم العالي والبحث العلمي رقم (40) لسنة 1988", "category": "تعليم عالي",
     "summary": "ينظم قطاع التعليم العالي والبحث العلمي ويحدد مهام الوزارة والجامعات ومراكز البحث.",
     "tags": ["تعليم عالي", "جامعات", "بحث علمي", "تدريسيين"],
     "articles": [
         {"num": "المادة 1", "text": "تتولى وزارة التعليم العالي والبحث العلمي رسم السياسة العامة للتعليم الجامعي والبحث العلمي"},
         {"num": "المادة 5", "text": "تتمتع الجامعات بالاستقلال الاداري والمالي والعلمي"},
         {"num": "المادة 14", "text": "يكون التعيين في الملاك التدريسي بناءً على شهادات علمية معتمدة وخبرة اكاديمية"},
         {"num": "المادة 22", "text": "تحدد الالقاب العلمية: مدرس مساعد، مدرس، استاذ مساعد، استاذ"},
     ]},
    {"id": 11, "title": "قانون هيئة النزاهة رقم (30) لسنة 2011", "category": "نزاهة ومكافحة فساد",
     "summary": "ينظم عمل هيئة النزاهة المستقلة لمكافحة الفساد الاداري والمالي ومتابعة الذمم المالية.",
     "tags": ["نزاهة", "فساد", "كشف ذمة", "رقابة"],
     "articles": [
         {"num": "المادة 1", "text": "هيئة النزاهة هيئة مستقلة تخضع لرقابة مجلس النواب"},
         {"num": "المادة 8", "text": "تتولى الهيئة مهمة منع ومكافحة الفساد والتحقيق في قضاياه"},
         {"num": "المادة 17", "text": "يلتزم الموظفون بتقديم كشف ذمتهم المالية خلال مدة محددة"},
         {"num": "المادة 23", "text": "عدم تقديم كشف الذمة المالية يعد مخالفة ادارية توجب المسائلة"},
     ]},
    {"id": 12, "title": "قانون حق الحصول على المعلومة رقم (18) لسنة 2013", "category": "حقوق ومعلومات",
     "summary": "يكفل حق المواطنين في الاطلاع على المعلومات الرسمية لدى الدوائر الحكومية.",
     "tags": ["حق المعلومة", "شفافية", "وصول"],
     "articles": [
         {"num": "المادة 2", "text": "لكل شخص الحق في الحصول على المعلومات المتوفرة لدى الجهات المشمولة بأحكام هذا القانون"},
         {"num": "المادة 7", "text": "يقدم الطلب كتابياً ويتم الرد خلال مدة لا تتجاوز 15 يوم عمل"},
         {"num": "المادة 12", "text": "يجوز رفض الطلب في حالات محددة كالامن الوطني او الخصوصية الشخصية"},
     ]},
    {"id": 13, "title": "قانون المحافظات غير المنتظمة في اقليم رقم (21) لسنة 2008", "category": "ادارة محلية",
     "summary": "ينظم صلاحيات المحافظات والمجالس المحلية وعلاقتها بالحكومة الاتحادية.",
     "tags": ["محافظات", "مجلس محلي", "لامركزية", "محافظ"],
     "articles": [
         {"num": "المادة 2", "text": "المحافظة وحدة ادارية تتمتع بالشخصية المعنوية والاستقلال المالي والاداري"},
         {"num": "المادة 7", "text": "يتكون مجلس المحافظة من اعضاء يُنتخبون بالاقتراع السري العام"},
         {"num": "المادة 24", "text": "للمحافظ صلاحيات تنفيذية تشمل ادارة شؤون المحافظة والاشراف على الدوائر"},
     ]},
    {"id": 14, "title": "قانون البلديات رقم (165) لسنة 1964", "category": "ادارة محلية",
     "summary": "ينظم عمل البلديات وصلاحياتها في الخدمات البلدية والتخطيط العمراني.",
     "tags": ["بلديات", "خدمات", "تخطيط عمراني"],
     "articles": [
         {"num": "المادة 1", "text": "البلدية مؤسسة عامة تتمتع بالشخصية المعنوية"},
         {"num": "المادة 9", "text": "تتولى البلدية تقديم الخدمات البلدية ضمن حدودها الادارية"},
         {"num": "المادة 15", "text": "يجوز للبلدية فرض رسوم مقابل الخدمات التي تقدمها وفق الانظمة"},
     ]},
    {"id": 15, "title": "قانون الاحزاب السياسية رقم (36) لسنة 2015", "category": "احزاب وسياسة",
     "summary": "ينظم تأسيس الاحزاب السياسية وحقوقها والتزاماتها القانونية.",
     "tags": ["احزاب", "سياسة", "انتخابات"],
     "articles": [
         {"num": "المادة 3", "text": "تأسيس الاحزاب السياسية حق يكفله الدستور"},
         {"num": "المادة 8", "text": "يشترط في مؤسسي الحزب ان يكونوا عراقيين ومن غير المحكومين بجرائم مخلة بالشرف"},
         {"num": "المادة 25", "text": "يحظر على الحزب اقامة تنظيمات عسكرية او شبه عسكرية"},
     ]},
    {"id": 16, "title": "قانون الضمان الاجتماعي للعمال رقم (39) لسنة 1971", "category": "ضمان اجتماعي",
     "summary": "ينظم حقوق العمال في الضمان الاجتماعي من اصابات عمل وتقاعد ومنافع اخرى.",
     "tags": ["ضمان اجتماعي", "عمال", "اصابات", "تقاعد"],
     "articles": [
         {"num": "المادة 2", "text": "يسري هذا القانون على العمال في القطاع الخاص والمختلط والتعاوني"},
         {"num": "المادة 8", "text": "يستحق العامل التعويض عن اصابات العمل وامراض المهنة"},
         {"num": "المادة 15", "text": "تحتسب مساهمة صاحب العمل والعامل في صندوق الضمان الاجتماعي"},
     ]},
    {"id": 17, "title": "قانون منع الاتجار بالبشر رقم (28) لسنة 2012", "category": "حقوق انسان",
     "summary": "يجرم الاتجار بالبشر ويحمي الضحايا ويفرض عقوبات صارمة على مرتكبيه.",
     "tags": ["اتجار بشر", "حقوق انسان", "عقوبات جنائية"],
     "articles": [
         {"num": "المادة 1", "text": "الاتجار بالبشر هو استدراج او نقل او ايواء اشخاص بقصد استغلالهم"},
         {"num": "المادة 6", "text": "يعاقب بالسجن المؤبد كل من ارتكب جريمة الاتجار بالبشر"},
         {"num": "المادة 10", "text": "توفر الدولة الحماية والمأوى للضحايا وتقدم لهم المساعدة القانونية"},
     ]},
    {"id": 18, "title": "قانون مفوضية حقوق الانسان رقم (53) لسنة 2008", "category": "حقوق انسان",
     "summary": "ينشئ المفوضية العليا لحقوق الانسان لحماية الحقوق والحريات العامة.",
     "tags": ["حقوق انسان", "حريات", "مفوضية"],
     "articles": [
         {"num": "المادة 1", "text": "تؤسس المفوضية العليا لحقوق الانسان كهيئة مستقلة"},
         {"num": "المادة 5", "text": "تتولى المفوضية تلقي الشكاوى المتعلقة بانتهاكات حقوق الانسان"},
         {"num": "المادة 9", "text": "للمفوضية صلاحية التحقيق في الشكاوى واصدار التوصيات للجهات المعنية"},
     ]},
    {"id": 19, "title": "قانون رعاية القاصرين رقم (78) لسنة 1980", "category": "رعاية",
     "summary": "ينظم رعاية القاصرين وادارة اموالهم وحقوقهم القانونية.",
     "tags": ["قاصر", "ولاية", "وصاية", "رعاية"],
     "articles": [
         {"num": "المادة 3", "text": "القاصر هو من لم يبلغ الثامنة عشرة من العمر"},
         {"num": "المادة 12", "text": "تكون الولاية على مال القاصر للاب ثم للجد لاب ثم للوصي"},
         {"num": "المادة 28", "text": "لا يجوز التصرف في اموال القاصر الا بإذن المحكمة المختصة"},
     ]},
    {"id": 20, "title": "قانون التنظيم النقابي للعمال رقم (52) لسنة 1987", "category": "عمال ونقابات",
     "summary": "ينظم حق العمال في تشكيل النقابات المهنية والانضمام اليها والدفاع عن حقوقهم.",
     "tags": ["نقابات", "عمال", "حقوق عمالية"],
     "articles": [
         {"num": "المادة 1", "text": "للعمال حق تشكيل النقابات والانضمام اليها بحرية"},
         {"num": "المادة 7", "text": "تتولى النقابة الدفاع عن حقوق العمال ومصالحهم المشروعة"},
         {"num": "المادة 15", "text": "يحق للنقابة التفاوض الجماعي نيابة عن اعضائها"},
     ]},
]

CATEGORIES = sorted(list(set(law["category"] for law in LAWS_DB)))

# === Legal Terms Glossary ===
LEGAL_TERMS = [
    {"term": "العقوبة الانضباطية", "definition": "جزاء ينزل بالموظف الذي يخالف واجبات وظيفته، ويتدرج من لفت النظر الى العزل.", "category": "انضباط", "related_law": "قانون الانضباط 14/1991"},
    {"term": "لفت النظر", "definition": "اخف العقوبات الانضباطية، تنبيه للموظف كتابياً دون اثر مالي.", "category": "انضباط", "related_law": "قانون الانضباط 14/1991 - المادة 8"},
    {"term": "الانذار", "definition": "عقوبة انضباطية تلي لفت النظر، تحذير رسمي للموظف.", "category": "انضباط", "related_law": "قانون الانضباط 14/1991 - المادة 8"},
    {"term": "قطع الراتب", "definition": "عقوبة انضباطية بخصم جزء من الراتب لمدة محددة.", "category": "انضباط", "related_law": "قانون الانضباط 14/1991 - المادة 8"},
    {"term": "التوبيخ", "definition": "عقوبة انضباطية رسمية تسجل في ملف الموظف.", "category": "انضباط", "related_law": "قانون الانضباط 14/1991 - المادة 8"},
    {"term": "الفصل", "definition": "انهاء خدمة الموظف بسبب مخالفات جسيمة، مع احتفاظه ببعض الحقوق.", "category": "انضباط", "related_law": "قانون الانضباط 14/1991"},
    {"term": "العزل", "definition": "اشد العقوبات الانضباطية، انهاء الخدمة مع حرمان من الحقوق التقاعدية كلياً او جزئياً.", "category": "انضباط", "related_law": "قانون الانضباط 14/1991 - المادة 8"},
    {"term": "الموظف", "definition": "كل شخص عهدت اليه وظيفة دائمة في ملاك الدولة او القطاع العام.", "category": "عام", "related_law": "قانون الخدمة المدنية 24/1960"},
    {"term": "التعيين", "definition": "اجراء اداري لاسناد وظيفة لشخص يستوفي الشروط القانونية.", "category": "خدمة مدنية", "related_law": "قانون الخدمة المدنية 24/1960 - المادة 7"},
    {"term": "التجربة", "definition": "مدة سنة واحدة من تاريخ المباشرة يُختبر فيها الموظف قبل تثبيته.", "category": "خدمة مدنية", "related_law": "قانون الخدمة المدنية 24/1960 - المادة 12"},
    {"term": "الاجازة الاعتيادية", "definition": "اجازة سنوية مدفوعة الراتب مدتها 30 يوم.", "category": "خدمة مدنية", "related_law": "قانون الخدمة المدنية 24/1960 - المادة 29"},
    {"term": "الاحالة على التقاعد", "definition": "انهاء خدمة الموظف عند بلوغه السن القانونية او اكمال مدة الخدمة.", "category": "تقاعد", "related_law": "قانون التقاعد الموحد 9/2014"},
    {"term": "الراتب التقاعدي", "definition": "راتب شهري يستحقه الموظف بعد احالته على التقاعد، يحسب على اساس 2.5% من معدل الراتب عن كل سنة خدمة.", "category": "تقاعد", "related_law": "قانون التقاعد 9/2014 - المادة 10"},
    {"term": "مكافأة نهاية الخدمة", "definition": "مبلغ يُمنح للموظف عند انتهاء خدمته دون استحقاقه للراتب التقاعدي.", "category": "تقاعد", "related_law": "قانون التقاعد 9/2014"},
    {"term": "الملاك", "definition": "جدول تنظيمي يبين عدد ونوع الوظائف الدائمة في الدائرة ودرجاتها.", "category": "ملاك", "related_law": "قانون الملاك 25/1960 - المادة 1"},
    {"term": "الدرجة الوظيفية", "definition": "مرتبة وظيفية تحدد بها مسؤوليات الموظف وراتبه، وتقسم الى 10 درجات.", "category": "ملاك", "related_law": "قانون الملاك 25/1960 - المادة 3"},
    {"term": "العلاوة السنوية", "definition": "زيادة سنوية في الراتب تُمنح للموظف.", "category": "ملاك", "related_law": "قانون الملاك 25/1960 - المادة 5"},
    {"term": "المخصصات", "definition": "مبالغ تُضاف الى الراتب الاسمي مثل مخصصات الشهادة والخدمة والمنصب.", "category": "رواتب", "related_law": "قانون الرواتب 22/2008 - المادة 2"},
    {"term": "القرار الاداري", "definition": "اجراء قانوني صادر عن جهة ادارية بارادتها المنفردة يرتب اثراً قانونياً.", "category": "قضاء اداري", "related_law": "قانون مجلس شورى الدولة 65/1979"},
    {"term": "الطعن الاداري", "definition": "وسيلة قانونية للاعتراض على القرار الاداري خلال 60 يوم من تاريخ التبليغ.", "category": "قضاء اداري", "related_law": "قانون مجلس شورى الدولة 65/1979 - المادة 12"},
    {"term": "مجلس شورى الدولة", "definition": "هيئة قضائية ادارية مستقلة تنظر في الطعون الادارية والفتاوى القانونية.", "category": "قضاء اداري", "related_law": "قانون مجلس شورى الدولة 65/1979 - المادة 1"},
    {"term": "الرشوة", "definition": "جريمة وظيفية بقبول الموظف هدية او منفعة مقابل اداء عمل او الامتناع عنه.", "category": "جرائم وظيفية", "related_law": "قانون العقوبات 111/1969 - المادة 307"},
    {"term": "الاختلاس", "definition": "استيلاء الموظف على مال عام اوكل اليه بحكم وظيفته.", "category": "جرائم وظيفية", "related_law": "قانون العقوبات 111/1969 - المادة 315"},
    {"term": "التزوير", "definition": "تغيير الحقيقة في محرر رسمي بقصد الاضرار.", "category": "جرائم وظيفية", "related_law": "قانون العقوبات 111/1969"},
    {"term": "استغلال الوظيفة", "definition": "استخدام الموظف لسلطته او نفوذه لتحقيق مصلحة شخصية.", "category": "جرائم وظيفية", "related_law": "قانون العقوبات 111/1969 - المادة 331"},
    {"term": "التحقيق الاداري", "definition": "اجراء قانوني للتحري عن مخالفة منسوبة للموظف قبل فرض العقوبة.", "category": "انضباط", "related_law": "قانون الانضباط 14/1991 - المادة 4"},
    {"term": "اللجنة الانضباطية", "definition": "لجنة تشكل في كل دائرة برئاسة موظف لا تقل درجته عن مدير للنظر في المخالفات.", "category": "انضباط", "related_law": "قانون الانضباط 14/1991 - المادة 15"},
    {"term": "الاستقالة", "definition": "انهاء الموظف لخدمته بارادته وفق اجراءات قانونية محددة.", "category": "خدمة مدنية", "related_law": "قانون الخدمة المدنية 24/1960 - المادة 36"},
    {"term": "النقل", "definition": "تغيير مكان عمل الموظف داخل او خارج دائرته وفق الاحتياج.", "category": "خدمة مدنية", "related_law": "قانون الخدمة المدنية 24/1960"},
    {"term": "الترقية", "definition": "انتقال الموظف الى درجة وظيفية اعلى وفق الاستحقاق والاقدمية.", "category": "خدمة مدنية", "related_law": "قانون الخدمة المدنية 24/1960"},
]

TERMS_CATEGORIES = sorted(list(set(t["category"] for t in LEGAL_TERMS)))

# === Legal Case Studies ===
CASE_STUDIES = [
    {
        "id": 1,
        "title": "موظف تعرض للفصل بسبب غياب متكرر",
        "category": "انضباط",
        "scenario": "موظف في دائرة حكومية تغيّب عن العمل لمدة 10 ايام متتالية دون عذر مشروع. شُكل له تحقيق اداري وتقرر فصله من الخدمة.",
        "question": "هل يحق للموظف الطعن بقرار الفصل؟ وما هي الاجراءات الصحيحة؟",
        "analysis": "وفقاً للمادة (4) من قانون انضباط موظفي الدولة رقم 14 لسنة 1991، لا يجوز فرض عقوبة انضباطية الا بعد اجراء تحقيق. وحسب المادة (10) من نفس القانون، للموظف حق الاعتراض على العقوبة المفروضة خلال 30 يوم من تاريخ التبليغ.",
        "legal_basis": ["قانون الانضباط 14/1991 - المادة 4", "قانون الانضباط 14/1991 - المادة 10"],
        "conclusion": "يحق للموظف الطعن بقرار الفصل امام لجنة الاعتراض خلال 30 يوم، ثم اللجوء الى مجلس شورى الدولة اذا رُفض اعتراضه خلال 60 يوم.",
    },
    {
        "id": 2,
        "title": "احتساب الراتب التقاعدي لموظف اكمل 25 سنة خدمة",
        "category": "تقاعد",
        "scenario": "موظف في كلية جامعية اكمل 25 سنة خدمة حكومية ومعدل راتبه السنوي في السنوات الاخيرة 1,200,000 دينار. طلب الاحالة على التقاعد.",
        "question": "كيف يُحتسب راتبه التقاعدي الشهري؟",
        "analysis": "وفقاً للمادة (10) من قانون التقاعد الموحد رقم 9 لسنة 2014، يحتسب الراتب التقاعدي على اساس 2.5% من معدل الراتب عن كل سنة خدمة.",
        "legal_basis": ["قانون التقاعد 9/2014 - المادة 9", "قانون التقاعد 9/2014 - المادة 10"],
        "conclusion": "الحساب: 25 × 2.5% = 62.5% من معدل الراتب\n1,200,000 × 62.5% = 750,000 دينار شهرياً كراتب تقاعدي.",
    },
    {
        "id": 3,
        "title": "طعن بقرار نقل موظف لدائرة اخرى",
        "category": "قضاء اداري",
        "scenario": "صدر قرار بنقل موظف من دائرته في بغداد الى محافظة اخرى دون اخذ موافقته ودون مبرر وظيفي. الموظف يعاني من ظروف عائلية.",
        "question": "ما الاجراءات القانونية المتاحة للموظف للطعن بالقرار؟",
        "analysis": "القرار الاداري يخضع للطعن وفقاً للمادة (12) من قانون مجلس شورى الدولة رقم 65 لسنة 1979، خلال 60 يوم من تاريخ التبليغ.",
        "legal_basis": ["قانون مجلس شورى الدولة 65/1979 - المادة 7", "قانون مجلس شورى الدولة 65/1979 - المادة 12"],
        "conclusion": "1. تقديم تظلم اداري للجهة المصدرة للقرار\n2. في حال رفض التظلم، يحق للموظف الطعن امام محكمة قضاء الموظفين خلال 60 يوم\n3. تقديم الاسباب الموضوعية والمستندات الداعمة.",
    },
    {
        "id": 4,
        "title": "موظف رفض اعادة تعيينه بعد العفو العام",
        "category": "جرائم وظيفية",
        "scenario": "موظف صدر بحقه حكم بالسجن بسبب جريمة اختلاس وفُصل من الخدمة. بعد صدور قانون عفو عام، طلب اعادة تعيينه فرُفض طلبه.",
        "question": "هل يحق لمن شمله العفو العام العودة الى وظيفته الحكومية؟",
        "analysis": "وفقاً للمادة (315) من قانون العقوبات رقم 111 لسنة 1969، من يُحكم عليه بالاختلاس يُعزل وجوباً من الخدمة. العفو العام يسقط العقوبة الجزائية لكن لا يُسقط العزل الاداري الذي هو عقوبة تبعية.",
        "legal_basis": ["قانون العقوبات 111/1969 - المادة 315", "قانون الانضباط 14/1991"],
        "conclusion": "العفو العام لا يعيد الموظف المعزول الى وظيفته تلقائياً، لأن العزل عقوبة ادارية مستقلة عن العقوبة الجزائية. يحتاج قرار خاص باعادة الاعتبار.",
    },
    {
        "id": 5,
        "title": "استحقاق مخصصات الشهادة العليا",
        "category": "رواتب",
        "scenario": "موظف حصل على شهادة الماجستير اثناء الخدمة وقدم طلباً لاحتساب مخصصات الشهادة العليا.",
        "question": "متى يستحق الموظف مخصصات الشهادة العليا؟",
        "analysis": "وفقاً للمادة (8) من قانون رواتب موظفي الدولة رقم 22 لسنة 2008، يستحق الموظف مخصصات الشهادة حسب التحصيل الدراسي من تاريخ حصوله عليها وتقديم المستندات الرسمية.",
        "legal_basis": ["قانون الرواتب 22/2008 - المادة 8"],
        "conclusion": "تُصرف المخصصات من تاريخ تقديم الطلب المرفق بالشهادة المصدقة، وتعتبر جزءاً من الراتب الكلي.",
    },
    {
        "id": 6,
        "title": "موظف طلب اجازة اعتيادية مرفوضة",
        "category": "خدمة مدنية",
        "scenario": "موظف قدم طلب اجازة اعتيادية لمدة 20 يوم لم يستخدم اي اجازة خلال السنة، لكن المدير رفض الطلب بحجة الحاجة الى العمل.",
        "question": "هل يحق للمدير رفض طلب الاجازة الاعتيادية؟",
        "analysis": "وفقاً للمادة (29) من قانون الخدمة المدنية رقم 24 لسنة 1960، يستحق الموظف اجازة اعتيادية براتب تام مدتها 30 يوم في السنة. الاجازة حق للموظف، لكن يمكن تأجيلها لظروف العمل لا رفضها.",
        "legal_basis": ["قانون الخدمة المدنية 24/1960 - المادة 29"],
        "conclusion": "لا يحق للمدير رفض الاجازة الاعتيادية بل يمكنه تأجيلها فقط لظروف العمل الملحة. يمكن للموظف تقديم اعتراض الى الجهة الاعلى.",
    },
]

CASE_CATEGORIES = sorted(list(set(c["category"] for c in CASE_STUDIES)))

def build_laws_context():
    ctx = ""
    for law in LAWS_DB:
        ctx += f"\n=== {law['title']} ===\nالتصنيف: {law['category']}\n{law['summary']}\n"
        for art in law["articles"]:
            ctx += f"  {art['num']}: {art['text']}\n"
    return ctx

LAWS_CONTEXT = build_laws_context()

# === File Processing ===
def extract_text_pdf(file_bytes):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(p.extract_text() or "" for p in reader.pages).strip()
    except Exception as e:
        return f"خطأ PDF: {e}"

def extract_text_docx(file_bytes):
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                r = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if r: parts.append(r)
        return "\n".join(parts)
    except Exception as e:
        return f"خطأ DOCX: {e}"

def process_file(uploaded_file):
    fb = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return {"type": "text", "content": extract_text_pdf(fb), "fmt": "PDF"}
    elif name.endswith((".docx", ".doc")):
        return {"type": "text", "content": extract_text_docx(fb), "fmt": "DOCX"}
    elif name.endswith(".txt"):
        return {"type": "text", "content": fb.decode("utf-8", errors="ignore"), "fmt": "TXT"}
    elif name.endswith((".png", ".jpg", ".jpeg", ".webp")):
        mt = "image/png" if name.endswith(".png") else ("image/webp" if name.endswith(".webp") else "image/jpeg")
        return {"type": "image", "content": base64.b64encode(fb).decode(), "media_type": mt, "fmt": "IMAGE"}
    return {"type": "error", "content": "نوع غير مدعوم. المدعوم: PDF, DOCX, TXT, PNG, JPG"}

# === OpenRouter API ===
OPENROUTER_TEXT_MODEL = "openai/gpt-oss-120b:free"
OPENROUTER_VISION_MODEL = "nvidia/nemotron-nano-12b-v2-vl:free"

def call_openrouter(api_key, system, user_content, use_vision=False):
    """user_content: string for text, or list of dicts for multimodal"""
    import urllib.request
    model = OPENROUTER_VISION_MODEL if use_vision else OPENROUTER_TEXT_MODEL
    body = {
        "model": model,
        "max_tokens": 4000,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user_content}
        ]
    }
    payload = json.dumps(body).encode("utf-8")
    req = urllib.request.Request(
        "https://openrouter.ai/api/v1/chat/completions", data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
            "HTTP-Referer": "https://legal-platform.local",
            "X-Title": "Iraqi Legal Platform",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            r = json.loads(resp.read().decode("utf-8"))
            return {"ok": True, "text": r["choices"][0]["message"]["content"]}
    except Exception as e:
        err_msg = str(e)
        if hasattr(e, 'read'):
            try:
                err_msg = e.read().decode("utf-8", errors="ignore")
            except:
                pass
        return {"ok": False, "text": err_msg}

def analyze_document_ai(api_key, doc, question=""):
    system = f"""You are an experienced Iraqi administrative legal consultant. Analyze the document and provide a complete, specific legal analysis. You MUST respond in Arabic only.

Iraqi Laws Database:
{LAWS_CONTEXT}

REQUIRED response structure:

## 1. ملخص الوثيقة
Summarize the specific document content in 3-5 sentences. Show you understood the actual content.

## 2. التحليل القانوني
Provide a DIRECT analysis of this SPECIFIC document. Explain what the legal issues are and how they apply to THIS case. Do not just list laws - analyze the specific situation step by step.

## 3. السند القانوني
For each applicable law: cite exact article number and text, then explain how it applies to THIS specific document.

## 4. التوصيات القانونية
Provide specific, actionable recommendations for what the person should do next.

## 5. ملاحظات
Any additional important legal considerations.

CRITICAL RULES:
- NEVER just list laws without connecting them to the document's specific content
- ALWAYS provide direct analysis of the actual scenario
- Be specific, not generic
- Write in formal Arabic
- No emojis
- Do NOT include any disclaimer"""

    if doc["type"] == "image":
        content = [
            {"type": "image_url", "image_url": {"url": f"data:{doc['media_type']};base64,{doc['content']}"}},
            {"type": "text", "text": f"حلل هذه الوثيقة وحدد السند القانوني.{chr(10) + 'سؤال: ' + question if question else ''}"}
        ]
        return call_openrouter(api_key, system, content, use_vision=True)
    user_text = f"حلل الوثيقة التالية وحدد السند القانوني:\n---\n{doc['content']}\n---{chr(10) + 'سؤال: ' + question if question else ''}"
    return call_openrouter(api_key, system, user_text)

def analyze_locally(text, question=""):
    combined = (text + " " + (question or "")).lower()
    kw_map = {
        "انضباط": [0], "عقوب": [0], "فصل": [0], "عزل": [0], "توبيخ": [0], "تحقيق": [0],
        "تعيين": [1], "ترقي": [1], "نقل": [1], "اجاز": [1], "مباشرة": [1], "استقال": [1],
        "تقاعد": [2], "معاش": [2], "متقاعد": [2],
        "ملاك": [3], "درج": [3, 6], "علاو": [3, 6],
        "حق": [4], "التزام": [4], "واجب": [4], "نزاه": [4],
        "شورى": [5], "طعن": [5, 8], "قضاء اداري": [5],
        "راتب": [6], "مخصص": [6],
        "رشوة": [7], "اختلاس": [7], "تزوير": [7], "استغلال": [7],
        "مرافع": [8], "دعوى": [8], "استئناف": [8], "تمييز": [8],
        "موظف": [0, 1, 4], "قرار اداري": [5],
    }
    matched = set()
    for kw, idx in kw_map.items():
        if kw in combined:
            matched.update(idx)
    if not matched:
        matched = {0, 1, 4}

    r = "## تحليل محلي\n\n### السند القانوني المحتمل\n\n"
    for i in sorted(matched):
        if i < len(LAWS_DB):
            law = LAWS_DB[i]
            r += f"**{law['title']}**\n{law['summary']}\n\n"
            for a in law["articles"]:
                r += f"- **{a['num']}**: {a['text']}\n"
            r += "\n---\n\n"
    return r

def chat_to_pdf(messages, title="محادثة المستشار القانوني"):
    """Export chat messages to PDF with proper Arabic support using fpdf2"""
    try:
        from fpdf import FPDF
        import arabic_reshaper
        from bidi.algorithm import get_display
        import re

        def ar(text):
            try:
                reshaped = arabic_reshaper.reshape(str(text))
                return get_display(reshaped)
            except Exception:
                return str(text)

        pdf = FPDF(orientation='P', unit='mm', format='A4')
        pdf.set_auto_page_break(auto=True, margin=18)

        font_paths = [
            ("C:/Windows/Fonts/tahoma.ttf", "C:/Windows/Fonts/tahomabd.ttf"),
            ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
        ]
        font_registered = False
        for regular, bold in font_paths:
            if os.path.exists(regular):
                try:
                    pdf.add_font("Arabic", "", regular)
                    pdf.add_font("Arabic", "B", bold if os.path.exists(bold) else regular)
                    font_registered = True
                    break
                except Exception:
                    continue
        if not font_registered:
            return None

        pdf.add_page()

        # Header
        pdf.set_fill_color(139, 26, 26)
        pdf.rect(0, 0, pdf.w, 22, 'F')
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Arabic", "B", 14)
        pdf.set_xy(10, 5)
        pdf.cell(pdf.w - 20, 12, ar("منصة القانون الاداري - جامعة ديالى"), align="R")
        pdf.ln(22)

        # Title and date
        pdf.set_font("Arabic", "B", 11)
        pdf.set_text_color(139, 26, 26)
        pdf.set_x(10)
        pdf.cell(pdf.w - 20, 8, ar(title), align="R", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Arabic", "", 8)
        pdf.set_text_color(120, 120, 120)
        pdf.set_x(10)
        pdf.cell(pdf.w - 20, 5, ar(f"التاريخ: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"), align="R", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(139, 26, 26)
        pdf.set_line_width(0.5)
        pdf.line(10, pdf.get_y() + 2, pdf.w - 10, pdf.get_y() + 2)
        pdf.ln(6)

        for msg in messages:
            content = msg.get("content", "")
            if not content:
                continue
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
            clean = clean.replace("**", "").replace("##", "").replace("---", "")

            role_name = "المستخدم" if msg.get("role") == "user" else "المستشار القانوني"
            color = (139, 26, 26) if msg.get("role") == "user" else (100, 50, 50)

            pdf.set_font("Arabic", "B", 10)
            pdf.set_text_color(*color)
            pdf.set_x(10)
            pdf.cell(pdf.w - 20, 7, ar(role_name + ":"), align="R", new_x="LMARGIN", new_y="NEXT")

            pdf.set_font("Arabic", "", 10)
            pdf.set_text_color(50, 50, 50)
            for para in clean.split('\n'):
                para = para.strip()
                if para:
                    pdf.set_x(10)
                    pdf.multi_cell(w=pdf.w - 20, h=6, text=ar(para), align="R")
                else:
                    pdf.ln(3)
            pdf.ln(4)
            pdf.set_draw_color(200, 200, 200)
            pdf.set_line_width(0.2)
            pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
            pdf.ln(4)

        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return None

def chat_response_stream(api_key, question, history):
    """Generator that yields text chunks as they arrive (streaming)"""
    if not api_key:
        yield local_chat(question)
        return
    system = f"""You are an experienced Iraqi administrative legal consultant specialized in Iraqi administrative law. You MUST respond in Arabic only, in clear and professional language.

Iraqi Laws Database:
{LAWS_CONTEXT}

CRITICAL INSTRUCTIONS - You MUST follow this response structure for every question:

## 1. فهم الحالة (Case Understanding)
Summarize the specific situation/scenario in 2-3 sentences to show you understood it.

## 2. التحليل القانوني (Legal Analysis)
Analyze the specific scenario directly. Explain what laws apply and WHY they apply to THIS particular case. Discuss the legal reasoning step by step.

## 3. السند القانوني (Legal Basis)
Cite the specific articles from Iraqi laws that apply, with the exact article number and text. Explain how each article relates to the case.

## 4. الخلاصة والتوصيات (Conclusion & Recommendations)
Give a clear, direct answer to what the person should do. Be specific and actionable.

RULES:
- NEVER just list laws without analyzing the specific case
- ALWAYS connect the legal articles to the actual scenario presented
- Be direct and specific in your analysis
- Write in formal Arabic
- No emojis
- Do NOT include any disclaimer about educational purposes"""
    msgs = [{"role": "system", "content": system}]
    for m in history[-6:]:
        msgs.append({"role": m["role"], "content": m["content"]})
    msgs.append({"role": "user", "content": question})
    import urllib.request
    body = {"model": OPENROUTER_TEXT_MODEL, "max_tokens": 2000, "messages": msgs, "stream": True}
    payload = json.dumps(body).encode("utf-8")
    req = urllib.request.Request("https://openrouter.ai/api/v1/chat/completions", data=payload,
        headers={
            "Content-Type": "application/json",
            "Accept": "text/event-stream",
            "Authorization": f"Bearer {api_key}",
            "HTTP-Referer": "https://legal-platform.local",
            "X-Title": "Iraqi Legal Platform",
            "User-Agent": "Mozilla/5.0"
        }, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=180) as resp:
            buffer = ""
            while True:
                line = resp.readline()
                if not line:
                    break
                try:
                    line_str = line.decode("utf-8").strip()
                except Exception:
                    continue
                if not line_str or line_str.startswith(":"):
                    continue
                if line_str.startswith("data: "):
                    data = line_str[6:]
                    if data == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data)
                        delta = chunk.get("choices", [{}])[0].get("delta", {}).get("content", "")
                        if delta:
                            yield delta
                    except json.JSONDecodeError:
                        continue
    except Exception as e:
        yield f"\n\n**خطأ:** {str(e)}"

def chat_response(api_key, question, history):
    if api_key:
        system = f"""You are an experienced Iraqi administrative legal consultant specialized in Iraqi administrative law. You MUST respond in Arabic only, in clear and professional language.

Iraqi Laws Database:
{LAWS_CONTEXT}

CRITICAL INSTRUCTIONS - You MUST follow this response structure for every question:

## 1. فهم الحالة (Case Understanding)
Summarize the specific situation/scenario in 2-3 sentences to show you understood it.

## 2. التحليل القانوني (Legal Analysis)
Analyze the specific scenario directly. Explain what laws apply and WHY they apply to THIS particular case. Discuss the legal reasoning step by step.

## 3. السند القانوني (Legal Basis)
Cite the specific articles from Iraqi laws that apply, with the exact article number and text. Explain how each article relates to the case.

## 4. الخلاصة والتوصيات (Conclusion & Recommendations)
Give a clear, direct answer to what the person should do. Be specific and actionable.

RULES:
- NEVER just list laws without analyzing the specific case
- ALWAYS connect the legal articles to the actual scenario presented
- Be direct and specific in your analysis
- Write in formal Arabic
- No emojis
- Do NOT include any disclaimer about educational purposes"""
        msgs = [{"role": "system", "content": system}]
        for m in history[-6:]:
            msgs.append({"role": m["role"], "content": m["content"]})
        msgs.append({"role": "user", "content": question})
        import urllib.request
        body = {"model": OPENROUTER_TEXT_MODEL, "max_tokens": 2000, "messages": msgs}
        payload = json.dumps(body).encode("utf-8")
        req = urllib.request.Request("https://openrouter.ai/api/v1/chat/completions", data=payload,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
                "HTTP-Referer": "https://legal-platform.local",
                "X-Title": "Iraqi Legal Platform",
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            }, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=180) as resp:
                r = json.loads(resp.read().decode("utf-8"))
                return r["choices"][0]["message"]["content"]
        except Exception as e:
            err_msg = str(e)
            if hasattr(e, 'read'):
                try:
                    err_msg = e.read().decode("utf-8", errors="ignore")
                except:
                    pass
            return f"**خطأ في الاتصال بـ OpenRouter API:**\n\n{err_msg}\n\n---\n**الرد المحلي البديل:**\n\n" + local_chat(question)
    return local_chat(question)

def local_chat(q):
    q = q.lower()
    kw = {"انضباط": [0], "عقوب": [0,7], "فصل": [0], "تعيين": [1], "ترقي": [1], "اجاز": [1],
          "تقاعد": [2], "ملاك": [3], "حق": [4], "شورى": [5], "طعن": [5], "راتب": [6],
          "رشوة": [7], "تزوير": [7], "اختلاس": [7], "استغلال": [7], "جريم": [7], "موظف": [0,1,4]}
    m = set()
    for k, v in kw.items():
        if k in q: m.update(v)
    if not m: m = {0, 1, 2}
    r = ""
    for i in sorted(m):
        if i < len(LAWS_DB):
            law = LAWS_DB[i]
            r += f"**{law['title']}**\n{law['summary']}\n\n"
            for a in law["articles"][:3]:
                r += f"- {a['num']}: {a['text']}\n"
            r += "\n---\n"
    return r

# === Sidebar ===
with st.sidebar:
    # === Brand Header (Hero Card) ===
    st.markdown("""
    <div class="sb-hero">
        <div class="sb-hero-logo">ق</div>
        <h1>منصة القانون الاداري</h1>
        <p>جامعة ديالى - رئاسة الجامعة</p>
    </div>
    """, unsafe_allow_html=True)

    # === [1] Navigation Menu ===
    st.markdown('<div class="sb-label"><span class="sb-label-bar"></span>القائمة الرئيسية</div>', unsafe_allow_html=True)
    nav_options = [
        "🏠  الرئيسية",
        "💬  المستشار القانوني",
        "📄  تحليل الوثائق",
        "📚  قاعدة القوانين",
        "📖  قاموس المصطلحات",
        "⚖️  دراسات الحالة",
        "📋  ادارة الاستشارات"
    ]
    if st.session_state.admin_logged_in:
        nav_options.append("🛡️  لوحة الادمن")

    page = st.radio("nav", nav_options, label_visibility="collapsed")
    page = page.split("  ", 1)[-1] if "  " in page else page

    # === [2] API Key (silent load) ===
    api_key = ""
    try:
        if hasattr(st, 'secrets') and "openrouter_api_key" in st.secrets:
            api_key = str(st.secrets["openrouter_api_key"]).strip()
    except Exception:
        pass
    if not api_key:
        api_key = os.environ.get("OPENROUTER_API_KEY", "").strip()

    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    # === [3] Quick Stats (compact) ===
    cons_s = st.session_state.data.get("consultations", [])
    anl_s = st.session_state.data.get("analyses", [])
    ta_s = sum(len(l["articles"]) for l in LAWS_DB)

    st.markdown(f"""
    <div class="sb-label"><span class="sb-label-bar"></span>نظرة سريعة</div>
    <div class="sb-stats-grid">
        <div class="sb-stat">
            <div class="sb-stat-num">{len(LAWS_DB)}</div>
            <div class="sb-stat-lbl">قانون</div>
        </div>
        <div class="sb-stat">
            <div class="sb-stat-num">{ta_s}</div>
            <div class="sb-stat-lbl">مادة</div>
        </div>
        <div class="sb-stat">
            <div class="sb-stat-num">{len(anl_s)}</div>
            <div class="sb-stat-lbl">تحليل</div>
        </div>
        <div class="sb-stat">
            <div class="sb-stat-num">{len(cons_s)}</div>
            <div class="sb-stat-lbl">استشارة</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    # === [5] Settings & Help (collapsed) ===
    st.markdown('<div class="sb-label"><span class="sb-label-bar"></span>الاعدادات</div>', unsafe_allow_html=True)

    with st.expander("🔤  حجم الخط"):
        font_size_choice = st.select_slider(
            "font_size",
            options=["صغير", "متوسط", "كبير", "كبير جداً"],
            value=st.session_state.font_size,
            label_visibility="collapsed",
            key="font_size_slider"
        )
        st.session_state.font_size = font_size_choice

    with st.expander("ℹ️  عن المنصة"):
        st.markdown("""
**منصة القانون الاداري** - مساعد ذكي يقدم:

- استشارات قانونية فورية
- تحليل الوثائق الرسمية
- قاعدة قوانين شاملة
- قاموس مصطلحات ودراسات حالة
        """)

    with st.expander("📘  دليل الاستخدام"):
        st.markdown("""
**1.** اختر القسم من القائمة

**2.** اكتب سؤالك في المستشار القانوني

**3.** ارفع وثيقة لتحليلها

**4.** تصفح القوانين والمصطلحات
        """)

    with st.expander("📱  تثبيت على الهاتف"):
        st.markdown("""
**للاندرويد (Chrome):**

1. افتح الموقع في Chrome
2. اضغط على القائمة (⋮)
3. اختر **"Add to Home screen"**
4. اضغط **Add**

**للايفون (Safari):**

1. افتح الموقع في Safari
2. اضغط زر المشاركة (□↑)
3. اختر **"Add to Home Screen"**
4. اضغط **Add**

بعدها سيظهر كتطبيق على شاشتك الرئيسية تماماً مثل أي تطبيق.
        """)

    # === [6] Designer Credits ===
    st.markdown("""
    <div class="sb-credits">
        <div class="sb-credits-line"></div>
        <p class="sb-credits-label">تصميم وتطوير</p>
        <div class="sb-credits-name">
            <div class="sb-credits-title">أ.م. علي حسين فاضل</div>
            <div class="sb-credits-sub">مصمم المنصة</div>
        </div>
        <div class="sb-credits-team">
            <div class="sb-team-icon">👥</div>
            <div class="sb-team-text">
                <div class="sb-team-title">كادر وحدة تكنولوجيا المعلومات</div>
            </div>
        </div>
    </div>

    <div class="sb-footer">
        <div class="sb-footer-line"></div>
        <p class="sb-footer-uni">جامعة ديالى</p>
        <p class="sb-footer-col">رئاسة الجامعة</p>
        <p class="sb-footer-year">© 2026</p>
    </div>
    """, unsafe_allow_html=True)

# === PAGES ===

# --- Home ---
if page == "الرئيسية":
    # Welcome voice - download MP3 server-side and embed as base64 (avoids CORS)
    @st.cache_data
    def get_welcome_audio_b64():
        try:
            import urllib.request, urllib.parse
            # Split into chunks (Google TTS has ~200 char limit per request)
            chunks = [
                "أهلا وسهلا بكم في منصة القانون الاداري لجامعة ديالى.",
                "هذه المنصة تقدم خدمات الاستشارات القانونية الادارية بشكل ذكي.",
                "تستطيع تحليل الوثائق الرسمية وتحديد السند القانوني المناسب.",
                "وتصفح قاعدة بيانات القوانين الادارية العراقية.",
                "والتحدث مع المستشار القانوني الذكي للاجابة على استفساراتك.",
                "نتمنى لكم تجربة مفيدة."
            ]
            all_bytes = b""
            for chunk in chunks:
                url = f"https://translate.google.com/translate_tts?ie=UTF-8&q={urllib.parse.quote(chunk)}&tl=ar&client=tw-ob"
                req = urllib.request.Request(url, headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
                })
                with urllib.request.urlopen(req, timeout=15) as resp:
                    all_bytes += resp.read()
            return base64.b64encode(all_bytes).decode("ascii")
        except Exception as e:
            return None

    audio_b64 = get_welcome_audio_b64()

    if audio_b64:
        components.html(f"""
        <!DOCTYPE html>
        <html>
        <head><meta charset="utf-8"></head>
        <body style="margin:0;padding:0;background:transparent;">
        <audio id="welcomeAudio" autoplay preload="auto" src="data:audio/mpeg;base64,{audio_b64}"></audio>
        <script>
        (function() {{
            const audio = document.getElementById('welcomeAudio');
            audio.volume = 1.0;
            let played = false;

            function play() {{
                if (played) return;
                audio.currentTime = 0;
                const p = audio.play();
                if (p !== undefined) {{
                    p.then(function() {{
                        played = true;
                        removeListeners();
                    }}).catch(function(err) {{}});
                }}
            }}

            function removeListeners() {{
                ['click','mousemove','keydown','scroll','touchstart','mousedown','pointerdown'].forEach(function(ev) {{
                    window.parent.document.removeEventListener(ev, play, true);
                    document.removeEventListener(ev, play, true);
                }});
            }}

            // Try immediate play
            play();

            // Also try on any user interaction (parent document since iframe)
            ['click','mousemove','keydown','scroll','touchstart','mousedown','pointerdown'].forEach(function(ev) {{
                try {{
                    window.parent.document.addEventListener(ev, play, {{capture: true, once: false}});
                }} catch(e) {{}}
                document.addEventListener(ev, play, {{capture: true, once: false}});
            }});

            // Also retry every 200ms for first 3 seconds
            let retries = 0;
            const interval = setInterval(function() {{
                if (played || retries > 15) {{
                    clearInterval(interval);
                    return;
                }}
                retries++;
                play();
            }}, 200);
        }})();
        </script>
        </body>
        </html>
        """, height=0)

    # Ensure there's a current conversation
    def get_current_conv():
        if not st.session_state.conversations:
            new_conv = {
                "id": new_conversation_id(),
                "title": "محادثة جديدة",
                "messages": [],
                "created": datetime.datetime.now().isoformat(),
                "updated": datetime.datetime.now().isoformat()
            }
            st.session_state.conversations.insert(0, new_conv)
            st.session_state.current_conv_id = new_conv["id"]
            save_conversations(st.session_state.conversations)
        if not st.session_state.current_conv_id:
            st.session_state.current_conv_id = st.session_state.conversations[0]["id"]
        for conv in st.session_state.conversations:
            if conv["id"] == st.session_state.current_conv_id:
                return conv
        # Fallback: use first
        st.session_state.current_conv_id = st.session_state.conversations[0]["id"]
        return st.session_state.conversations[0]

    current_conv = get_current_conv()

    # === Conversations Manager Bar ===
    conv_c1, conv_c2, conv_c3 = st.columns([5, 1, 1])
    with conv_c1:
        conv_options = {c["id"]: c["title"] for c in st.session_state.conversations}
        conv_ids = list(conv_options.keys())
        current_idx = conv_ids.index(st.session_state.current_conv_id) if st.session_state.current_conv_id in conv_ids else 0
        selected_conv_id = st.selectbox(
            "المحادثات",
            options=conv_ids,
            format_func=lambda cid: conv_options.get(cid, "محادثة")[:50],
            index=current_idx,
            label_visibility="collapsed",
            key="conv_selector"
        )
        if selected_conv_id != st.session_state.current_conv_id:
            st.session_state.current_conv_id = selected_conv_id
            st.rerun()
    with conv_c2:
        if st.button("➕ جديدة", use_container_width=True, key="new_conv_btn"):
            new_conv = {
                "id": new_conversation_id(),
                "title": "محادثة جديدة",
                "messages": [],
                "created": datetime.datetime.now().isoformat(),
                "updated": datetime.datetime.now().isoformat()
            }
            st.session_state.conversations.insert(0, new_conv)
            st.session_state.current_conv_id = new_conv["id"]
            save_conversations(st.session_state.conversations)
            st.rerun()
    with conv_c3:
        if len(st.session_state.conversations) > 1:
            if st.button("🗑️ حذف", use_container_width=True, key="del_conv_btn"):
                st.session_state.conversations = [c for c in st.session_state.conversations if c["id"] != st.session_state.current_conv_id]
                st.session_state.current_conv_id = st.session_state.conversations[0]["id"]
                save_conversations(st.session_state.conversations)
                st.rerun()

    # Bind chat_messages to current conversation for compatibility
    st.session_state.chat_messages = current_conv["messages"]

    def md_to_html(text):
        """Convert simple markdown to HTML"""
        import re
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        # Headers
        text = re.sub(r'^### (.+)$', r'<h4 style="color:#8B1A1A;margin:0.5rem 0;">\1</h4>', text, flags=re.MULTILINE)
        text = re.sub(r'^## (.+)$', r'<h3 style="color:#8B1A1A;margin:0.5rem 0;">\1</h3>', text, flags=re.MULTILINE)
        # List items
        text = re.sub(r'^- (.+)$', r'• \1', text, flags=re.MULTILINE)
        # Horizontal rule
        text = re.sub(r'^---+$', r'<hr style="border:none;border-top:1px solid #e0d0d0;margin:0.5rem 0;">', text, flags=re.MULTILINE)
        # Newlines
        text = text.replace("\n", "<br>")
        return text

    # Build entire chat as ONE HTML block
    msgs_html = """<div class="msg-row bot">
<div class="msg-avatar bot-av">ق</div>
<div class="msg-content">
<div class="msg-bubble bot-bubble">مرحبا بك! انا المستشار القانوني الذكي لجامعة ديالى<br>كيف يمكنني مساعدتك اليوم؟</div>
<div class="msg-time">المستشار القانوني</div>
</div>
</div>
"""

    for m in st.session_state.chat_messages:
        if m["role"] == "user":
            content = md_to_html(m["content"])
            msgs_html += f'<div class="msg-row user"><div class="msg-content user-content"><div class="msg-bubble user-bubble">{content}</div><div class="msg-time user-time">انت</div></div><div class="msg-avatar user-av">أ</div></div>\n'
        else:
            content = md_to_html(m["content"])
            msgs_html += f'<div class="msg-row bot"><div class="msg-avatar bot-av">ق</div><div class="msg-content"><div class="msg-bubble bot-bubble">{content}</div><div class="msg-time">المستشار القانوني</div></div></div>\n'

    # Render full chat using components.html (full HTML support, no markdown interference)
    chat_full_html = f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8">
<style>
@import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800&display=swap');
* {{ font-family: 'Tajawal', sans-serif; box-sizing: border-box; }}
body {{ margin: 0; padding: 0; background: transparent; direction: rtl; }}
.chat-conversation {{
    background: linear-gradient(rgba(255,255,255,0.95), rgba(255,255,255,0.95)),
                repeating-linear-gradient(45deg, #fdf0f0 0, #fdf0f0 2px, #fff 2px, #fff 12px);
    border: 1px solid #e0e0e0; border-radius: 16px;
    padding: 1rem; overflow-y: auto;
    display: flex; flex-direction: column; gap: 0.7rem;
}}
.chat-conversation::-webkit-scrollbar {{ width: 8px; }}
.chat-conversation::-webkit-scrollbar-track {{ background: #f5f5f5; }}
.chat-conversation::-webkit-scrollbar-thumb {{ background: #c0a0a0; border-radius: 4px; }}
.msg-row {{ display: flex; gap: 10px; direction: rtl; align-items: flex-end; }}
.msg-row.bot {{ justify-content: flex-start; }}
.msg-row.user {{ justify-content: flex-start; flex-direction: row-reverse; }}
.msg-avatar {{
    width: 42px; height: 42px; flex-shrink: 0; border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-size: 1.1rem; font-weight: 800; color: #fff;
}}
.bot-av {{ background: linear-gradient(135deg, #8B1A1A 0%, #a52a2a 100%); box-shadow: 0 3px 10px rgba(139,26,26,0.3); }}
.user-av {{ background: linear-gradient(135deg, #555 0%, #333 100%); box-shadow: 0 3px 10px rgba(0,0,0,0.2); }}
.msg-content {{ max-width: 75%; display: flex; flex-direction: column; }}
.user-content {{ align-items: flex-end; }}
.msg-bubble {{
    padding: 0.9rem 1.2rem; line-height: 1.8; font-size: 0.95rem;
    word-wrap: break-word; direction: rtl; text-align: right;
}}
.bot-bubble {{
    background: #fff; border-radius: 18px 18px 18px 4px;
    color: #333; border: 1px solid #f0e0e0;
    box-shadow: 0 2px 8px rgba(139,26,26,0.06);
}}
.bot-bubble strong {{ color: #8B1A1A; }}
.user-bubble {{
    background: linear-gradient(135deg, #8B1A1A 0%, #a52a2a 100%);
    color: #fff; border-radius: 18px 18px 4px 18px;
    box-shadow: 0 2px 10px rgba(139,26,26,0.25);
}}
.user-bubble strong {{ color: #fff; }}
.msg-time {{
    font-size: 0.7rem; color: #999; margin: 4px 8px 0 8px;
    direction: rtl; text-align: right; font-weight: 600;
}}
.user-time {{ text-align: left; color: #8B1A1A; }}
</style>
</head>
<body>
<div class="chat-conversation">
{msgs_html}
</div>
<script>
window.scrollTo(0, document.body.scrollHeight);
const conv = document.querySelector('.chat-conversation');
if (conv) conv.scrollTop = conv.scrollHeight;
</script>
</body>
</html>"""

    # Calculate height based on messages count
    base_height = 200
    msg_height = 120 * (len(st.session_state.chat_messages) + 1)
    chat_height = min(max(base_height, msg_height), 600)
    components.html(chat_full_html, height=chat_height, scrolling=True)

    # Suggested questions
    if len(st.session_state.chat_messages) == 0:
        st.markdown('<div class="suggestions-title">اقتراحات للبدء:</div>', unsafe_allow_html=True)
        qcs = st.columns(2)
        qqs = ["ما هي العقوبات الانضباطية على الموظف؟", "ما شروط التقاعد للموظف العراقي؟", "ما هي حقوق الموظف عند الفصل؟", "كيف اطعن بقرار اداري؟"]
        home_q = None
        for i, q in enumerate(qqs):
            with qcs[i % 2]:
                if st.button(q, key=f"hq{i}", use_container_width=True): home_q = q
    else:
        home_q = None

    # Input bar - using form so Enter submits
    with st.form(key="chat_form", clear_on_submit=True):
        ic1, ic2, ic3 = st.columns([7, 1, 1])
        with ic1:
            home_input = st.text_input("msg", placeholder="اكتب رسالتك هنا واضغط Enter...", key="home_chat_input", label_visibility="collapsed")
        with ic2:
            send_home = st.form_submit_button("➤", use_container_width=True, type="primary", help="ارسال")
        with ic3:
            clear_clicked = st.form_submit_button("🗑", use_container_width=True, help="مسح المحادثة")

    if clear_clicked:
        st.session_state.chat_messages = []
        st.rerun()

    home_question = home_q or (home_input if send_home and home_input else None)
    if home_question:
        current_conv["messages"].append({"role": "user", "content": home_question})
        # Auto-title from first message
        if current_conv["title"] == "محادثة جديدة":
            current_conv["title"] = home_question[:40] + ("..." if len(home_question) > 40 else "")

        # Show thinking indicator
        with st.spinner("المستشار القانوني يحلل الحالة..."):
            full_response = chat_response(api_key, home_question, current_conv["messages"][:-1])

        if not full_response or not full_response.strip():
            full_response = local_chat(home_question)

        current_conv["messages"].append({"role": "assistant", "content": full_response})
        current_conv["updated"] = datetime.datetime.now().isoformat()
        save_conversations(st.session_state.conversations)
        st.rerun()

    # === Export PDF Button (ALWAYS visible) ===
    st.markdown('<div style="margin-top:1rem;"></div>', unsafe_allow_html=True)
    has_msgs = len(current_conv.get("messages", [])) > 0
    if has_msgs:
        try:
            pdf_data = chat_to_pdf(current_conv["messages"], current_conv.get("title", "محادثة"))
        except Exception:
            pdf_data = None

        if pdf_data and len(pdf_data) > 100:
            st.download_button(
                label="تصدير المحادثة كـ PDF",
                data=pdf_data,
                file_name=f"محادثة_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="pdf_export_main"
            )
        else:
            # PDF generation failed - offer text download as fallback
            text_export = ""
            for m in current_conv["messages"]:
                role = "المستخدم" if m["role"] == "user" else "المستشار"
                text_export += f"\n{'='*40}\n{role}:\n{m['content']}\n"
            st.download_button(
                label="تصدير المحادثة كملف نصي",
                data=text_export,
                file_name=f"محادثة_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain",
                use_container_width=True,
                key="txt_export_main"
            )
    else:
        st.button("تصدير المحادثة كـ PDF", disabled=True, use_container_width=True, key="pdf_disabled_main")

# --- Document Analysis ---
elif page == "تحليل الوثائق":
    st.markdown('<div class="breadcrumb">القائمة &lt; تحليل الوثائق</div>', unsafe_allow_html=True)
    st.markdown('<div class="page-title">تحليل الوثائق والكتب الرسمية</div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="how-box">
        <p class="how-title">كيف يعمل:</p>
        <p style="color:#555;margin:0;font-size:0.9rem;line-height:2;">
            1. ارفع الوثيقة (PDF, Word, صورة, نص) او اكتب نص المشكلة<br>
            2. النظام يقرأ ويحلل الموضوع القانوني<br>
            3. يحدد القوانين والمواد ذات الصلة<br>
            4. يقدم السند القانوني والتوصيات
        </p>
    </div>
    """, unsafe_allow_html=True)

    tab_up, tab_txt, tab_hist = st.tabs(["رفع وثيقة", "كتابة نص المشكلة", "سجل التحليلات"])

    with tab_up:
        uploaded = st.file_uploader("اختر الملف", type=["pdf","docx","doc","txt","png","jpg","jpeg","webp"])
        extra_q = st.text_area("سؤال اضافي (اختياري)", placeholder="مثال: ما السند القانوني للطعن بهذا القرار؟", height=80)

        if uploaded:
            st.markdown(f'<div class="file-info"><p class="name">{uploaded.name}</p><p class="size">{uploaded.size/1024:.1f} KB</p></div>', unsafe_allow_html=True)

        if st.button("تحليل الوثيقة", type="primary", use_container_width=True, disabled=not uploaded):
            with st.spinner("جاري قراءة الوثيقة..."):
                doc = process_file(uploaded)
            if doc["type"] == "error":
                st.error(doc["content"])
            else:
                with st.spinner("جاري التحليل القانوني..."):
                    if api_key:
                        res = analyze_document_ai(api_key, doc, extra_q)
                        txt = res["text"] if res["ok"] else analyze_locally(doc.get("content",""), extra_q)
                    elif doc["type"] == "text":
                        txt = analyze_locally(doc["content"], extra_q)
                    else:
                        st.warning("تحليل الصور يتطلب مفتاح OpenRouter API.")
                        txt = None

                if txt:
                    st.session_state.analysis_result = txt
                    rec = {"id": f"DOC-{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}",
                           "filename": uploaded.name, "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                           "question": extra_q or "تحليل عام", "result_preview": txt[:200], "full_result": txt}
                    st.session_state.data.setdefault("analyses", []).append(rec)
                    save_data(st.session_state.data)

        if st.session_state.analysis_result:
            st.markdown("---")
            st.markdown("### نتيجة التحليل القانوني")
            st.markdown(f'<div class="analysis-box"></div>', unsafe_allow_html=True)
            st.markdown(st.session_state.analysis_result)
            if st.button("مسح النتيجة"):
                st.session_state.analysis_result = None
                st.rerun()

    with tab_txt:
        problem = st.text_area("نص المشكلة القانونية", placeholder="مثال: صدر قرار بنقل موظف دون موافقته...", height=250)
        tq = st.text_area("سؤال محدد (اختياري)", placeholder="هل يحق الطعن؟", height=80, key="tq")

        if st.button("تحليل النص", type="primary", use_container_width=True, disabled=not problem):
            with st.spinner("جاري التحليل..."):
                doc = {"type": "text", "content": problem}
                if api_key:
                    res = analyze_document_ai(api_key, doc, tq)
                    txt = res["text"] if res["ok"] else analyze_locally(problem, tq)
                else:
                    txt = analyze_locally(problem, tq)

            st.session_state.data.setdefault("analyses", []).append({
                "id": f"TXT-{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}",
                "filename": "نص مكتوب", "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                "question": tq or "تحليل عام", "result_preview": txt[:200], "full_result": txt})
            save_data(st.session_state.data)
            st.markdown("---")
            st.markdown("### نتيجة التحليل")
            st.markdown(txt)

    with tab_hist:
        analyses = st.session_state.data.get("analyses", [])
        if not analyses:
            st.info("لا توجد تحليلات سابقة.")
        else:
            st.markdown(f"**عدد التحليلات: {len(analyses)}**")
            for a in reversed(analyses):
                st.markdown(f'<div class="consultation-card"><div style="display:flex;justify-content:space-between;flex-direction:row-reverse;"><h4 style="color:#8B1A1A;margin:0;">{a["filename"]}</h4><span style="color:#888;font-size:0.8rem;">{a["date"]}</span></div><p style="color:#666;font-size:0.85rem;">السؤال: {a["question"]}</p></div>', unsafe_allow_html=True)
                with st.expander(f"التحليل الكامل - {a['id']}"):
                    st.markdown(a["full_result"])

# --- Laws DB ---
elif page == "قاعدة القوانين":
    st.markdown('<div class="breadcrumb">القائمة &lt; قاعدة القوانين</div>', unsafe_allow_html=True)
    st.markdown('<div class="page-title">قاعدة بيانات القوانين</div>', unsafe_allow_html=True)

    def normalize_ar(text):
        """Normalize Arabic: remove diacritics, alef variants, ta marbuta"""
        import re
        t = re.sub(r'[\u064B-\u065F\u0670]', '', text)  # diacritics
        t = t.replace('أ', 'ا').replace('إ', 'ا').replace('آ', 'ا')
        t = t.replace('ة', 'ه').replace('ى', 'ي')
        return t.lower().strip()

    def smart_search_score(query, law):
        """Calculate relevance score for a law based on query"""
        if not query.strip():
            return 1
        q_norm = normalize_ar(query)
        q_words = [w for w in q_norm.split() if len(w) >= 2]
        if not q_words:
            return 0

        score = 0
        title_norm = normalize_ar(law["title"])
        summary_norm = normalize_ar(law["summary"])
        tags_norm = [normalize_ar(t) for t in law.get("tags", [])]
        articles_norm = [normalize_ar(a["text"]) for a in law.get("articles", [])]

        for word in q_words:
            if word in title_norm: score += 10
            if word in summary_norm: score += 5
            for tag in tags_norm:
                if word in tag: score += 3
            for art in articles_norm:
                if word in art: score += 2
            # Partial match
            if any(word in title_norm[i:i+len(word)+2] for i in range(0, max(1, len(title_norm)-len(word)))):
                score += 1
        return score

    st.markdown("""
    <div style="background:#fdf0f0;border-right:4px solid #8B1A1A;border-radius:10px;padding:0.8rem;margin-bottom:1rem;">
        <p style="color:#8B1A1A;margin:0;font-weight:700;">🔍 بحث ذكي</p>
        <p style="color:#666;margin:2px 0 0 0;font-size:0.85rem;">يبحث في العناوين والمواد والتاغات مع دعم تطابق الكلمات المشتقة</p>
    </div>
    """, unsafe_allow_html=True)

    cs, cf = st.columns([3, 1])
    with cs: sq = st.text_input("البحث", placeholder="مثال: عقوبة الموظف، تقاعد، طعن...")
    with cf: sc = st.selectbox("التصنيف", ["الكل"] + CATEGORIES)

    # Smart search with scoring
    results = []
    for law in LAWS_DB:
        if sc != "الكل" and law["category"] != sc:
            continue
        score = smart_search_score(sq, law) if sq else 1
        if score > 0:
            results.append((score, law))
    results.sort(key=lambda x: -x[0])

    if sq:
        st.markdown(f"**عدد النتائج: {len(results)} — مرتّبة حسب الصلة**")
    else:
        st.markdown(f"**عدد القوانين: {len(results)}**")

    for score, law in results:
        tags = "".join(f'<span class="law-tag">{t}</span>' for t in law.get("tags", []))
        relevance = ""
        if sq and score > 0:
            relevance = f'<span style="background:#e8f5e9;color:#2e7d32;padding:2px 8px;border-radius:12px;font-size:0.75rem;margin-right:8px;">صلة: {min(100, score*5)}%</span>'
        st.markdown(f'<div class="law-card"><div style="display:flex;justify-content:space-between;align-items:center;"><h4 style="margin:0;">{law["title"]}</h4>{relevance}</div><p>{law["summary"]}</p><div style="margin-top:8px;">{tags}</div></div>', unsafe_allow_html=True)
        with st.expander(f"المواد - {law['title'][:50]}..."):
            for a in law["articles"]:
                # Highlight matching words
                text = a['text']
                if sq:
                    for w in sq.split():
                        if len(w) >= 2:
                            text = text.replace(w, f"**{w}**")
                st.markdown(f"**{a['num']}**: {text}")

# --- Legal Terms Glossary ---
elif page == "قاموس المصطلحات":
    st.markdown('<div class="breadcrumb">القائمة &lt; قاموس المصطلحات</div>', unsafe_allow_html=True)
    st.markdown('<div class="page-title">قاموس المصطلحات القانونية</div>', unsafe_allow_html=True)

    tc1, tc2 = st.columns([3, 1])
    with tc1: term_search = st.text_input("ابحث عن مصطلح", placeholder="مثال: الانذار، العزل، التقاعد...")
    with tc2: term_cat = st.selectbox("التصنيف", ["الكل"] + TERMS_CATEGORIES, key="term_cat")

    def term_matches(term, query):
        if not query.strip(): return True
        q = query.strip()
        return q in term["term"] or q in term["definition"] or q in term.get("category", "")

    filtered_terms = [t for t in LEGAL_TERMS if term_matches(t, term_search) and (term_cat == "الكل" or t["category"] == term_cat)]
    st.markdown(f"**عدد المصطلحات: {len(filtered_terms)}**")

    for t in filtered_terms:
        st.markdown(f"""
        <div style="background:#fff;border:1px solid #e0e0e0;border-right:4px solid #8B1A1A;border-radius:10px;padding:1rem;margin-bottom:0.8rem;">
            <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.5rem;">
                <h4 style="color:#8B1A1A;margin:0;">{t["term"]}</h4>
                <span style="background:#fdf0f0;color:#8B1A1A;padding:3px 10px;border-radius:12px;font-size:0.75rem;">{t["category"]}</span>
            </div>
            <p style="color:#444;margin:0.3rem 0;line-height:1.8;">{t["definition"]}</p>
            <p style="color:#888;margin:0.5rem 0 0 0;font-size:0.8rem;"><strong>السند القانوني:</strong> {t.get("related_law", "-")}</p>
        </div>
        """, unsafe_allow_html=True)

# --- Case Studies ---
elif page == "دراسات الحالة":
    st.markdown('<div class="breadcrumb">القائمة &lt; دراسات الحالة</div>', unsafe_allow_html=True)
    st.markdown('<div class="page-title">دراسات الحالة القانونية</div>', unsafe_allow_html=True)

    st.markdown("""
    <div style="background:#fdf0f0;border-right:4px solid #8B1A1A;border-radius:10px;padding:0.8rem;margin-bottom:1rem;">
        <p style="color:#8B1A1A;margin:0;font-weight:700;">⚖️ قضايا واقعية</p>
        <p style="color:#666;margin:2px 0 0 0;font-size:0.85rem;">حالات قانونية حقيقية مع التحليل والسند القانوني والخلاصة</p>
    </div>
    """, unsafe_allow_html=True)

    cc1, cc2 = st.columns([3, 1])
    with cc1: case_search = st.text_input("ابحث في الحالات", placeholder="مثال: فصل، تقاعد، نقل...", key="case_search")
    with cc2: case_cat = st.selectbox("التصنيف", ["الكل"] + CASE_CATEGORIES, key="case_cat")

    filtered_cases = [c for c in CASE_STUDIES
                      if (case_cat == "الكل" or c["category"] == case_cat)
                      and (not case_search.strip() or case_search in c["title"] or case_search in c.get("scenario", ""))]

    st.markdown(f"**عدد الحالات: {len(filtered_cases)}**")

    for c in filtered_cases:
        with st.container():
            st.markdown(f"""
            <div style="background:#fff;border:1px solid #e0e0e0;border-right:4px solid #8B1A1A;border-radius:12px;padding:1.2rem;margin-bottom:1rem;">
                <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.5rem;">
                    <h4 style="color:#8B1A1A;margin:0;">📋 {c["title"]}</h4>
                    <span style="background:#fdf0f0;color:#8B1A1A;padding:3px 10px;border-radius:12px;font-size:0.75rem;">{c["category"]}</span>
                </div>
            </div>
            """, unsafe_allow_html=True)
            with st.expander(f"🔍 عرض تفاصيل الحالة - {c['title']}"):
                st.markdown(f"**📝 السيناريو:**\n\n{c['scenario']}")
                st.markdown(f"**❓ السؤال القانوني:**\n\n{c['question']}")
                st.markdown(f"**⚖️ التحليل القانوني:**\n\n{c['analysis']}")
                st.markdown("**📚 السند القانوني:**")
                for lb in c.get("legal_basis", []):
                    st.markdown(f"- {lb}")
                st.markdown(f"**✅ الخلاصة:**\n\n{c['conclusion']}")

# --- Chat ---
elif page == "المستشار القانوني":
    st.markdown('<div class="breadcrumb">القائمة &lt; المستشار القانوني</div>', unsafe_allow_html=True)
    st.markdown('<div class="page-title">المستشار القانوني الذكي</div>', unsafe_allow_html=True)

    st.markdown("**اسئلة سريعة:**")
    qcs = st.columns(4)
    qqs = ["ما هي العقوبات الانضباطية؟", "شروط التقاعد؟", "حقوق الموظف؟", "كيف اطعن بقرار اداري؟"]
    sq = None
    for i, c in enumerate(qcs):
        with c:
            if st.button(qqs[i], key=f"q{i}", use_container_width=True): sq = qqs[i]
    st.markdown("---")

    if not st.session_state.chat_messages:
        st.markdown('<div class="chat-ai">مرحبا! انا المستشار القانوني. اسألني عن القوانين الادارية العراقية.</div>', unsafe_allow_html=True)
    else:
        for m in st.session_state.chat_messages:
            cls = "chat-user" if m["role"] == "user" else "chat-ai"
            st.markdown(f'<div class="{cls}">{m["content"]}</div>', unsafe_allow_html=True)

    ui = st.text_input("سؤالك", placeholder="مثال: حقوق الموظف عند الفصل؟", key="ci")
    c1, c2 = st.columns(2)
    with c1: send = st.button("ارسال", use_container_width=True, type="primary")
    with c2:
        if st.button("مسح", use_container_width=True):
            st.session_state.chat_messages = []
            st.rerun()

    question = sq or (ui if send and ui else None)
    if question:
        st.session_state.chat_messages.append({"role": "user", "content": question})
        with st.spinner("جاري البحث..."):
            resp = chat_response(api_key, question, st.session_state.chat_messages[:-1])
        st.session_state.chat_messages.append({"role": "assistant", "content": resp})
        st.rerun()

# --- Consultations ---
elif page == "ادارة الاستشارات":
    st.markdown('<div class="breadcrumb">القائمة &lt; ادارة الاستشارات</div>', unsafe_allow_html=True)
    st.markdown('<div class="page-title">ادارة الاستشارات</div>', unsafe_allow_html=True)
    t1, t2, t3 = st.tabs(["الاستشارات", "جديدة", "احصائيات"])
    cons = st.session_state.data.get("consultations", [])

    with t1:
        if not cons:
            st.info("لا توجد استشارات.")
        else:
            fc1, fc2 = st.columns(2)
            with fc1: fs = st.selectbox("الحالة", ["الكل","جديدة","قيد المراجعة","تمت الاجابة","مغلقة"])
            with fc2: fca = st.selectbox("التصنيف", ["الكل"]+CATEGORIES, key="fcc")
            fl = cons
            if fs != "الكل": fl = [c for c in fl if c["status"] == fs]
            if fca != "الكل": fl = [c for c in fl if c["category"] == fca]
            for i, c in enumerate(reversed(fl)):
                scls = {"جديدة":"status-new","قيد المراجعة":"status-progress","تمت الاجابة":"status-done","مغلقة":"status-closed"}.get(c["status"],"status-new")
                st.markdown(f'<div class="consultation-card"><div style="display:flex;justify-content:space-between;align-items:center;flex-direction:row-reverse;"><h4 style="color:#8B1A1A;margin:0;">{c["title"]}</h4><span class="status-badge {scls}">{c["status"]}</span></div><p style="color:#888;font-size:0.8rem;">{c["id"]} | {c["category"]} | {c["date"]}</p><p style="color:#555;">{c["description"][:200]}</p></div>', unsafe_allow_html=True)
                ai = len(cons) - 1 - i
                with st.expander(f"تحديث - {c['id']}"):
                    ns = st.selectbox("حالة", ["جديدة","قيد المراجعة","تمت الاجابة","مغلقة"],
                        index=["جديدة","قيد المراجعة","تمت الاجابة","مغلقة"].index(c["status"]), key=f"s_{c['id']}")
                    nt = st.text_area("ملاحظات", value=c.get("notes",""), key=f"n_{c['id']}")
                    if st.button("حفظ", key=f"v_{c['id']}"):
                        st.session_state.data["consultations"][ai]["status"] = ns
                        st.session_state.data["consultations"][ai]["notes"] = nt
                        save_data(st.session_state.data)
                        st.success("تم"); st.rerun()

    with t2:
        ti = st.text_input("العنوان", placeholder="استفسار حول الترقية")
        ca = st.selectbox("التصنيف", CATEGORIES, key="nc")
        pr = st.selectbox("الاولوية", ["عادية","متوسطة","عاجلة"])
        de = st.text_area("التفاصيل", height=150)
        ap = st.text_input("الاسم (اختياري)")
        if st.button("تسجيل", type="primary", use_container_width=True):
            if not ti or not de: st.error("ملء العنوان والتفاصيل")
            else:
                cid = f"CON-{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
                st.session_state.data["consultations"].append({"id":cid,"title":ti,"category":ca,"priority":pr,"description":de,"applicant":ap,"status":"جديدة","date":datetime.datetime.now().strftime("%Y-%m-%d"),"notes":""})
                save_data(st.session_state.data)
                st.success(f"تم - {cid}"); st.rerun()

    with t3:
        if not cons: st.info("لا بيانات")
        else:
            total = len(cons)
            sc = st.columns(4)
            for col, (s, clr) in zip(sc, [("جديدة","#2e7d32"),("قيد المراجعة","#f57f17"),("تمت الاجابة","#1565c0"),("مغلقة","#c62828")]):
                with col:
                    cnt = sum(1 for c in cons if c["status"]==s)
                    st.markdown(f'<div class="stat-card"><h2 style="color:{clr};">{cnt}</h2><p>{s}</p></div>', unsafe_allow_html=True)

# --- Admin Panel ---
elif page == "لوحة الادمن":
    if not st.session_state.admin_logged_in:
        st.error("غير مخول للوصول")
        st.stop()

    st.markdown('<div class="breadcrumb">القائمة &lt; لوحة الادمن</div>', unsafe_allow_html=True)
    st.markdown('<div class="page-title">لوحة ادارة المناطق الجغرافية</div>', unsafe_allow_html=True)

    admin_tab1, admin_tab2, admin_tab3, admin_tab4 = st.tabs(["🗺️ الخريطة والتأشيرات", "📊 الاحصائيات", "📋 سجل الدخول", "⚙️ الاعدادات"])

    geo_log = load_geo_log()
    markers = st.session_state.geo_config.get("markers", [])
    polygon = st.session_state.geo_config.get("polygon", [])

    # --- Tab 1: Buildings Dropdown + Polygon Drawing ---
    with admin_tab1:
        st.markdown("### تحديد الحدود الجغرافية للمنصة")
        st.info("1️⃣ اختر المبنى من القائمة المنسدلة (رئاسة او تشكيلات)\n\n2️⃣ اضغط على الخريطة لاضافة نقاط المضلع\n\n3️⃣ احفظ المضلع لتطبيق الحدود على المنصة")

        # === Dropdown: select building ===
        dc1, dc2 = st.columns([3, 1])
        with dc1:
            building_options = ["-- اختر المبنى --"] + [f"{b['name']} ({b['type']})" for b in UOD_BUILDINGS]
            selected_idx = st.selectbox("المبنى الجامعي", range(len(building_options)), format_func=lambda i: building_options[i], key="bldg_select")
        with dc2:
            locate_clicked = st.button("🎯 تحديد الموقع", type="primary", use_container_width=True)

        selected_building = None
        if selected_idx > 0:
            selected_building = UOD_BUILDINGS[selected_idx - 1]
            st.markdown(f"""
            <div style="background:#fdf0f0;border:1px solid #e8c4c4;border-radius:10px;padding:0.8rem;margin:0.5rem 0;">
                <strong style="color:#8B1A1A;">📍 {selected_building['name']}</strong><br>
                <span style="color:#666;font-size:0.85rem;">النوع: {selected_building['type']} | الاحداثيات: {selected_building['lat']:.5f}, {selected_building['lon']:.5f}</span>
            </div>
            """, unsafe_allow_html=True)

        # Store selected building in session for map
        if locate_clicked and selected_building:
            st.session_state.selected_building = selected_building
            st.rerun()

        current_building = st.session_state.get("selected_building")

        # === Current Polygon Points ===
        st.markdown("---")
        st.markdown("#### 🔺 نقاط المضلع الحالية")

        if "polygon_draft" not in st.session_state:
            st.session_state.polygon_draft = list(polygon) if polygon else []

        draft = st.session_state.polygon_draft

        if not draft:
            st.warning("⚠️ لا توجد نقاط. اضغط على الخريطة لاضافة نقاط المضلع.")
        else:
            st.success(f"✓ عدد النقاط الحالية: {len(draft)}")
            for i, pt in enumerate(draft):
                pc1, pc2, pc3 = st.columns([0.5, 4, 1])
                with pc1: st.markdown(f"**{i+1}**")
                with pc2: st.text(f"{pt[0]:.6f}, {pt[1]:.6f}")
                with pc3:
                    if st.button("حذف", key=f"del_pt_{i}"):
                        st.session_state.polygon_draft.pop(i)
                        st.rerun()

        # === Manual point entry ===
        with st.expander("➕ اضافة نقطة يدوياً"):
            mpc1, mpc2, mpc3 = st.columns([2, 2, 1])
            with mpc1: manual_lat = st.number_input("Lat", value=33.7632, format="%.6f", key="manual_lat")
            with mpc2: manual_lon = st.number_input("Lon", value=44.6183, format="%.6f", key="manual_lon")
            with mpc3:
                if st.button("اضافة", key="add_manual_pt"):
                    st.session_state.polygon_draft.append([float(manual_lat), float(manual_lon)])
                    st.rerun()

        # === Action buttons ===
        st.markdown("---")
        ac1, ac2, ac3, ac4 = st.columns(4)
        with ac1:
            if st.button("💾 حفظ المضلع", type="primary", use_container_width=True):
                if len(st.session_state.polygon_draft) >= 3:
                    st.session_state.geo_config["polygon"] = list(st.session_state.polygon_draft)
                    save_geo_config(st.session_state.geo_config)
                    st.success(f"تم حفظ مضلع بـ {len(st.session_state.polygon_draft)} نقطة")
                    st.rerun()
                else:
                    st.error("يحتاج على الاقل 3 نقاط")
        with ac2:
            if st.button("🔄 توليد تلقائي (Convex Hull)", use_container_width=True):
                if len(st.session_state.polygon_draft) >= 3:
                    hull = convex_hull(st.session_state.polygon_draft)
                    st.session_state.polygon_draft = hull
                    st.success(f"تم توليد مغلف محدب بـ {len(hull)} نقطة")
                    st.rerun()
                else:
                    st.error("يحتاج 3 نقاط على الاقل")
        with ac3:
            if st.button("↩️ استرجاع المحفوظ", use_container_width=True):
                st.session_state.polygon_draft = list(polygon) if polygon else []
                st.rerun()
        with ac4:
            if st.button("🗑️ مسح الكل", use_container_width=True):
                st.session_state.polygon_draft = []
                st.rerun()

        st.markdown("---")
        st.markdown("#### 🗺️ الخريطة التفاعلية")
        st.caption("اضغط في اي مكان على الخريطة لاضافة نقطة مضلع | المبنى المحدد يظهر بأيقونة عنابية")

        # Build map data
        buildings_data = json.dumps(UOD_BUILDINGS)
        draft_polygon_data = json.dumps(st.session_state.polygon_draft)
        saved_polygon_data = json.dumps(polygon)
        selected_data = json.dumps(current_building) if current_building else "null"
        log_points = [[e["lat"], e["lon"], 1 if e["allowed"] else 0] for e in geo_log if e.get("lat") is not None]
        log_points_json = json.dumps(log_points[-300:])

        map_html = """<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css">
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<style>
body { margin: 0; padding: 0; font-family: 'Tajawal', sans-serif; }
#map { width: 100%; height: 560px; border-radius: 12px; cursor: crosshair; }
.info-box { padding: 10px; background: #fff; border-radius: 8px; margin-top: 8px; direction: rtl; text-align: right; font-size: 12px; display: flex; gap: 15px; flex-wrap: wrap; justify-content: center; border: 1px solid #e0e0e0; }
.info-box span { display: inline-flex; align-items: center; gap: 5px; }
.status-bar { padding: 10px; background: linear-gradient(135deg,#8B1A1A,#a52a2a); color: #fff; border-radius: 8px; margin-bottom: 8px; direction: rtl; text-align: center; font-size: 13px; font-weight: 700; }
</style>
</head>
<body>
<div class="status-bar" id="statusBar">اضغط على الخريطة لاضافة نقاط المضلع - النقاط الحالية: <span id="ptCount">__DRAFT_LEN__</span></div>
<div id="map"></div>
<div class="info-box">
<span>🏛️ مبنى مختار</span>
<span>🏢 مباني جامعة ديالى</span>
<span style="color:#8B1A1A;font-weight:700;">━ مضلع مسودة</span>
<span style="color:#22c55e;">━ مضلع محفوظ</span>
<span style="color:#22c55e;">●</span> <span>دخول ناجح</span>
<span style="color:#ef4444;">●</span> <span>دخول مرفوض</span>
</div>
<script>
const BUILDINGS = __BUILDINGS__;
const SELECTED = __SELECTED__;
const DRAFT_POLYGON = __DRAFT_POLYGON__;
const SAVED_POLYGON = __SAVED_POLYGON__;
const LOG_POINTS = __LOG_POINTS__;

// Initial map center
let centerLat = 33.7632, centerLon = 44.6183, zoom = 15;
if (SELECTED) {
    centerLat = SELECTED.lat;
    centerLon = SELECTED.lon;
    zoom = 17;
} else if (DRAFT_POLYGON.length > 0) {
    centerLat = DRAFT_POLYGON[0][0];
    centerLon = DRAFT_POLYGON[0][1];
    zoom = 16;
}

const map = L.map('map').setView([centerLat, centerLon], zoom);
L.tileLayer('https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png', {
    attribution: '© OpenStreetMap',
    maxZoom: 19
}).addTo(map);

// Selected building icon (big gold)
const selectedIcon = L.divIcon({
    html: '<div style="background:#8B1A1A;color:#fff;width:44px;height:44px;border-radius:50% 50% 50% 0;transform:rotate(-45deg);display:flex;align-items:center;justify-content:center;border:4px solid #FFD700;box-shadow:0 4px 15px rgba(0,0,0,0.5);animation:pulse 1.5s infinite;"><span style="transform:rotate(45deg);font-size:20px;">🏛️</span></div><style>@keyframes pulse{0%,100%{transform:scale(1);}50%{transform:scale(1.1);}}</style>',
    className: '',
    iconSize: [44, 44],
    iconAnchor: [22, 44]
});

// Regular building icon
const buildingIcon = L.divIcon({
    html: '<div style="background:#666;color:#fff;width:28px;height:28px;border-radius:50%;display:flex;align-items:center;justify-content:center;border:2px solid #fff;box-shadow:0 2px 6px rgba(0,0,0,0.3);font-size:12px;">🏢</div>',
    className: '',
    iconSize: [28, 28],
    iconAnchor: [14, 14]
});

// Draw all buildings
BUILDINGS.forEach(function(b) {
    const isSelected = SELECTED && SELECTED.id === b.id;
    const marker = L.marker([b.lat, b.lon], { icon: isSelected ? selectedIcon : buildingIcon }).addTo(map);
    marker.bindPopup('<div style="direction:rtl;text-align:right;font-family:Tajawal,sans-serif;"><b>' + b.name + '</b><br>النوع: ' + b.type + '<br><small>' + b.lat.toFixed(5) + ', ' + b.lon.toFixed(5) + '</small></div>');
});

// Draw saved polygon (green)
if (SAVED_POLYGON && SAVED_POLYGON.length >= 3) {
    L.polygon(SAVED_POLYGON, {
        color: '#22c55e',
        weight: 2,
        fillColor: '#22c55e',
        fillOpacity: 0.08,
        dashArray: '4,4'
    }).addTo(map).bindPopup('المضلع المحفوظ (' + SAVED_POLYGON.length + ' نقطة)');
}

// Draw draft polygon (maroon, more prominent)
let draftLayer = null;
let draftMarkers = [];
function drawDraft() {
    if (draftLayer) map.removeLayer(draftLayer);
    draftMarkers.forEach(function(m) { map.removeLayer(m); });
    draftMarkers = [];

    if (DRAFT_POLYGON.length >= 3) {
        draftLayer = L.polygon(DRAFT_POLYGON, {
            color: '#8B1A1A',
            weight: 3,
            fillColor: '#8B1A1A',
            fillOpacity: 0.15
        }).addTo(map);
    } else if (DRAFT_POLYGON.length === 2) {
        draftLayer = L.polyline(DRAFT_POLYGON, {
            color: '#8B1A1A',
            weight: 3,
            dashArray: '6,4'
        }).addTo(map);
    }

    // Add numbered markers for each point
    DRAFT_POLYGON.forEach(function(p, i) {
        const ptIcon = L.divIcon({
            html: '<div style="background:#8B1A1A;color:#fff;width:24px;height:24px;border-radius:50%;display:flex;align-items:center;justify-content:center;border:2px solid #fff;font-weight:800;font-size:11px;box-shadow:0 2px 5px rgba(0,0,0,0.3);">' + (i+1) + '</div>',
            className: '',
            iconSize: [24, 24],
            iconAnchor: [12, 12]
        });
        draftMarkers.push(L.marker([p[0], p[1]], { icon: ptIcon }).addTo(map));
    });
}
drawDraft();

// Draw log points
LOG_POINTS.forEach(function(p) {
    L.circleMarker([p[0], p[1]], {
        radius: 4,
        color: p[2] === 1 ? '#22c55e' : '#ef4444',
        fillColor: p[2] === 1 ? '#22c55e' : '#ef4444',
        fillOpacity: 0.6,
        weight: 1
    }).addTo(map);
});

// Click on map to add polygon point
map.on('click', function(e) {
    const lat = e.latlng.lat.toFixed(6);
    const lon = e.latlng.lng.toFixed(6);
    // Send point to parent via URL
    const parent = window.parent;
    const url = new URL(parent.location.href);
    url.searchParams.set('add_pt', lat + ',' + lon);
    parent.location.href = url.toString();
});
</script>
</body>
</html>"""
        map_html = (map_html
            .replace("__BUILDINGS__", buildings_data)
            .replace("__SELECTED__", selected_data)
            .replace("__DRAFT_POLYGON__", draft_polygon_data)
            .replace("__SAVED_POLYGON__", saved_polygon_data)
            .replace("__LOG_POINTS__", log_points_json)
            .replace("__DRAFT_LEN__", str(len(st.session_state.polygon_draft))))
        components.html(map_html, height=660)

    # --- Tab 2: Statistics ---
    with admin_tab2:
        st.markdown("### احصائيات الدخول بخوارزمية Haversine")

        total = len(geo_log)
        allowed_count = sum(1 for e in geo_log if e.get("allowed"))
        denied_count = total - allowed_count
        allowed_pct = (allowed_count / total * 100) if total > 0 else 0
        denied_pct = (denied_count / total * 100) if total > 0 else 0

        sc1, sc2, sc3, sc4 = st.columns(4)
        with sc1:
            st.markdown(f'<div class="stat-card"><h2>{total}</h2><p>اجمالي المحاولات</p></div>', unsafe_allow_html=True)
        with sc2:
            st.markdown(f'<div class="stat-card"><h2 style="color:#2e7d32;">{allowed_count}</h2><p>داخل الحدود</p></div>', unsafe_allow_html=True)
        with sc3:
            st.markdown(f'<div class="stat-card"><h2 style="color:#c62828;">{denied_count}</h2><p>خارج الحدود</p></div>', unsafe_allow_html=True)
        with sc4:
            st.markdown(f'<div class="stat-card"><h2>{allowed_pct:.1f}%</h2><p>نسبة النجاح</p></div>', unsafe_allow_html=True)

        if total > 0:
            st.markdown("### التوزيع المئوي")
            st.markdown(f"""
            <div style="background:#fff;border:1px solid #e0e0e0;border-radius:12px;padding:1.5rem;margin:1rem 0;">
                <div style="display:flex;justify-content:space-between;margin-bottom:0.5rem;">
                    <span style="color:#2e7d32;font-weight:700;">داخل الحدود: {allowed_pct:.1f}%</span>
                    <span style="color:#c62828;font-weight:700;">خارج الحدود: {denied_pct:.1f}%</span>
                </div>
                <div style="display:flex;height:30px;border-radius:15px;overflow:hidden;border:1px solid #ddd;">
                    <div style="width:{allowed_pct}%;background:linear-gradient(135deg,#22c55e,#16a34a);"></div>
                    <div style="width:{denied_pct}%;background:linear-gradient(135deg,#ef4444,#dc2626);"></div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Statistics per nearest marker
            st.markdown("### توزيع الدخول حسب اقرب مبنى جامعي")
            if markers and allowed_count > 0:
                marker_stats = {m["id"]: 0 for m in markers}
                for e in geo_log:
                    if e.get("allowed") and e.get("lat") is not None:
                        # Find nearest marker
                        nearest = min(markers, key=lambda m: haversine_km(e["lat"], e["lon"], m["lat"], m["lon"]))
                        marker_stats[nearest["id"]] = marker_stats.get(nearest["id"], 0) + 1

                for m in markers:
                    count = marker_stats.get(m["id"], 0)
                    pct = (count / allowed_count * 100) if allowed_count > 0 else 0
                    badge = " 🏛️" if m.get("is_main") else ""
                    st.markdown(f"""
                    <div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:1rem;margin:0.5rem 0;">
                        <div style="display:flex;justify-content:space-between;align-items:center;">
                            <strong style="color:#8B1A1A;">{m["name"]}{badge}</strong>
                            <span style="color:#666;">{count} دخول ({pct:.1f}%)</span>
                        </div>
                        <div style="background:#f0e0e0;height:8px;border-radius:4px;margin-top:8px;overflow:hidden;">
                            <div style="width:{pct}%;background:linear-gradient(135deg,#8B1A1A,#a52a2a);height:100%;"></div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

            # Distance distribution
            st.markdown("### متوسط المسافات للمحاولات المرفوضة")
            denied_distances = [e.get("distance_km", 0) for e in geo_log if not e.get("allowed") and e.get("distance_km") is not None]
            if denied_distances:
                avg_dist = sum(denied_distances) / len(denied_distances)
                max_dist = max(denied_distances)
                min_dist = min(denied_distances)
                dc1, dc2, dc3 = st.columns(3)
                with dc1: st.metric("المتوسط", f"{avg_dist:.2f} كم")
                with dc2: st.metric("الاقل", f"{min_dist:.2f} كم")
                with dc3: st.metric("الاقصى", f"{max_dist:.2f} كم")
        else:
            st.info("لا توجد محاولات دخول مسجلة بعد")

    # --- Tab 3: Access Log ---
    with admin_tab3:
        st.markdown("### سجل محاولات الدخول")
        if not geo_log:
            st.info("لا يوجد سجل")
        else:
            filter_opt = st.radio("تصفية", ["الكل", "داخل الحدود فقط", "خارج الحدود فقط"], horizontal=True)
            filtered = geo_log
            if filter_opt == "داخل الحدود فقط":
                filtered = [e for e in geo_log if e.get("allowed")]
            elif filter_opt == "خارج الحدود فقط":
                filtered = [e for e in geo_log if not e.get("allowed")]

            st.markdown(f"**عدد النتائج: {len(filtered)}**")
            for e in reversed(filtered[-50:]):
                status_color = "#2e7d32" if e.get("allowed") else "#c62828"
                status_text = "مسموح" if e.get("allowed") else "مرفوض"
                ts = e.get("timestamp", "").replace("T", " ")[:19]
                dist = e.get("distance_km", "N/A")
                lat = e.get("lat", "N/A")
                lon = e.get("lon", "N/A")
                st.markdown(f"""
                <div style="background:#fff;border:1px solid #e0e0e0;border-right:4px solid {status_color};border-radius:8px;padding:0.8rem;margin:0.4rem 0;">
                    <div style="display:flex;justify-content:space-between;">
                        <span style="color:{status_color};font-weight:700;">{status_text}</span>
                        <span style="color:#888;font-size:0.8rem;">{ts}</span>
                    </div>
                    <div style="color:#666;font-size:0.85rem;margin-top:4px;">
                        الاحداثيات: {lat}, {lon} | المسافة: {dist} كم
                    </div>
                </div>
                """, unsafe_allow_html=True)

            if st.button("مسح السجل بالكامل", type="secondary"):
                with open(GEO_LOG_FILE, "w", encoding="utf-8") as f:
                    json.dump([], f)
                st.success("تم مسح السجل")
                st.rerun()

    # --- Tab 4: Settings ---
    with admin_tab4:
        st.markdown("### اعدادات الادمن")
        new_pwd = st.text_input("كلمة مرور جديدة", type="password")
        if st.button("تحديث كلمة المرور"):
            if new_pwd:
                st.session_state.geo_config["admin_password"] = new_pwd
                save_geo_config(st.session_state.geo_config)
                st.success("تم تحديث كلمة المرور")
            else:
                st.error("يرجى ادخال كلمة مرور")

        st.markdown("---")
        st.markdown("### معلومات النظام")
        st.json({
            "markers_count": len(markers),
            "included_markers": sum(1 for m in markers if m.get("included")),
            "main_marker": next((m["name"] for m in markers if m.get("is_main")), "غير محدد"),
            "polygon_points": len(polygon),
            "total_log_entries": len(geo_log),
            "config_file": GEO_CONFIG_FILE,
            "log_file": GEO_LOG_FILE
        })

        if st.button("تسجيل خروج الادمن", type="primary"):
            st.session_state.admin_logged_in = False
            st.session_state.geo_verified = False
            st.rerun()
